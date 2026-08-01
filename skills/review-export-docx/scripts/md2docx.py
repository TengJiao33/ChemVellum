#!/usr/bin/env python3
"""
md2docx.py -- Convert a Markdown file to DOCX using review_template.docx styles.

Inline support:
  **bold**  *italic*  ***bold-italic***  `code`
  ^superscript^       alnum_subscript_   $math$  $$display math$$
  [@citation]  ->  [citation key]

Section-aware styling (auto-detected from headings OR bold-only paragraphs):
  Abstract / Keywords / Acknowledgments / References / Supporting Information

Default layout profile (professional_single):
  H1 title      : Times New Roman 18 pt
  Section heads : Times New Roman 13.5 / 11.5 / 10.5 pt
  Body          : Times New Roman 11 pt, justified, 1.15 spacing
  Captions      : Times New Roman 9 pt, left aligned
  References    : Times New Roman 9 pt with real Word numbering

Usage:
    python3 scripts/md2docx.py --input review.md --output review.docx
"""
from __future__ import annotations

import argparse
import os
import re
from copy import deepcopy  # noqa: F401
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml  # noqa: F401
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None

try:
    from latex2word import LatexToWordElement
    _LATEX_OK = True
except ImportError:
    _LATEX_OK = False

# ---------------------------------------------------------------------------
# Template style name map
# ---------------------------------------------------------------------------

_S: Dict[str, str] = {
    "title":        "Review Title",
    "author":       "Review Author",
    "address":      "Review Affiliation",
    "email":        "Review Affiliation",
    "abstract":     "Review Abstract",
    "keywords":     "Review Keywords",
    "body":         "Review Body",
    "body_lead":    "Review Body Lead",
    "list":         "Review List",
    "toc":          "Review TOC",
    "h1":           "Review Heading 1",
    "h2":           "Review Heading 2",
    "h3":           "Review Heading 3",
    "figure":       "Review Figure Caption",
    "figure_display":"Review Figure Display",
    "table_title":  "Review Table Caption",
    "table_header": "Review Table Header",
    "table_body":   "Review Table Body",
    "chart":        "Review Figure Caption",
    "scheme":       "Review Figure Caption",
    "references":   "Review Reference",
    "acks":         "Review Body",
    "supporting":   "Review Body",
    "footnote":     "Review Reference",
}

# ---------------------------------------------------------------------------
# Font spec -- used for math objects and deliberate character-style overrides.
# ---------------------------------------------------------------------------

_FONT_SPEC: Dict[str, Dict] = {
    "title":        {"font": "Times New Roman", "size": 18, "bold": True},
    "author":       {"font": "Times New Roman", "size": 12},
    "address":      {"font": "Times New Roman", "size": 10.5},
    "email":        {"font": "Times New Roman", "size": 10.5},
    "abstract":     {"font": "Times New Roman", "size": 10.5},
    "keywords":     {"font": "Times New Roman", "size": 10.5},
    "body":         {"font": "Times New Roman", "size": 11},
    "body_lead":    {"font": "Times New Roman", "size": 11},
    "list":         {"font": "Times New Roman", "size": 11},
    "toc":          {"font": "Times New Roman", "size": 10},
    "h1":           {"font": "Times New Roman", "size": 13.5, "bold": True},
    "h2":           {"font": "Times New Roman", "size": 11.5, "bold": True},
    "h3":           {"font": "Times New Roman", "size": 10.5, "bold": True, "italic": True},
    "h4":           {"font": "Times New Roman", "size": 10.5, "italic": True},
    "figure":       {"font": "Times New Roman", "size": 9},
    "table_title":  {"font": "Times New Roman", "size": 9},
    "table_header": {"font": "Times New Roman", "size": 9, "bold": True},
    "table_body":   {"font": "Times New Roman", "size": 9},
    "scheme":       {"font": "Times New Roman", "size": 9},
    "chart":        {"font": "Times New Roman", "size": 9},
    "references":   {"font": "Times New Roman", "size": 9},
    "acks":         {"font": "Times New Roman", "size": 11},
    "supporting":   {"font": "Times New Roman", "size": 11},
    "footnote":     {"font": "Times New Roman", "size": 9},
}

_DEFAULT_LAYOUT_PROFILE = "professional_single"
_LAYOUT_PROFILES = ("professional_single", "chemvellum_journal", "legacy_report")

_CHEMVELLUM = {
    "teal": "0F6B67",
    "ink": "1E2A2D",
    "muted": "647176",
    "pale": "EAF3F2",
    "abstract_pale": "E1EEEC",
    "row_pale": "F7FAFA",
    "rule": "C9D9D7",
    "page_width": 8.5,
    "page_height": 11.0,
    "top_margin": 0.62,
    "bottom_margin": 0.65,
    "side_margin": 0.65,
    "header_distance": 0.27,
    "footer_distance": 0.30,
    "content_width": 7.2,
    "content_width_dxa": 10368,
    "column_gap_dxa": 360,
}

# Heading level -> (para_style_key, font_spec_key)
_HEADING_FORMAT: Dict[int, Tuple[str, str]] = {
    1: ("title", "title"),
    2: ("h1",  "h1"),
    3: ("h2",  "h2"),
    4: ("h3",  "h3"),
    5: ("body",  "body"),
    6: ("body",  "body"),
}

_SECTION_CONTEXT: Dict[str, str] = {
    "abstract":               "abstract",
    "keywords":               "keywords",
    "key words":              "keywords",
    "acknowledgments":        "acks",
    "acknowledgements":       "acks",
    "supporting information": "supporting",
    "references":             "references",
    "reference":              "references",
}

_CAPTION_PATTERNS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r"^(figure|fig\.)\s*\d+", re.I), "figure"),
    (re.compile(r"^table\s*\d+",            re.I), "table_title"),
    (re.compile(r"^scheme\s*\d+",           re.I), "scheme"),
    (re.compile(r"^chart\s*\d+",            re.I), "chart"),
]


def _usable_page_width_inches(
    doc: Document,
    layout_profile: str = _DEFAULT_LAYOUT_PROFILE,
) -> float:
    section = doc.sections[0]
    width_emu = section.page_width - section.left_margin - section.right_margin
    width_inches = width_emu / 914400
    # The portable single-column profile keeps the previous conservative cap.
    if layout_profile != "chemvellum_journal":
        width_inches = min(6.2, width_inches)
    return max(1.0, width_inches)


def _bounded_figure_size(path: Path, max_width: float, max_height: float = 5.9) -> Tuple[float, Optional[float]]:
    """Bound figure height so Word can keep its caption on the same page."""
    if PILImage is None:
        return max_width, None
    try:
        with PILImage.open(path) as image:
            width_px, height_px = image.size
    except Exception:
        return max_width, None
    if width_px <= 0 or height_px <= 0:
        return max_width, None
    ratio = width_px / height_px
    width = min(max_width, max_height * ratio)
    return width, width / ratio


def _figure_bounds_for_caption(caption_key: str, page_width: float) -> Tuple[float, float]:
    """Use a compact displayed block for reaction schemes."""
    if caption_key == "scheme":
        return min(page_width, 5.25), 3.6
    return page_width, 5.9


def _set_style_font(
    style,
    font_name: str,
    size: float,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def _ensure_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _set_outline_level(style, level: int | None) -> None:
    ppr = style.element.get_or_add_pPr()
    existing = ppr.find(qn("w:outlineLvl"))
    if existing is not None:
        ppr.remove(existing)
    if level is not None:
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level))
        ppr.append(outline)


def _configure_style(
    doc: Document,
    name: str,
    *,
    size: float,
    bold: bool = False,
    italic: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.0,
    keep_with_next: bool = False,
    keep_together: bool = False,
    first_line_indent: float = 0,
    left_indent: float = 0,
    outline_level: int | None = None,
    font_name: str = "Times New Roman",
    color: str | None = None,
):
    style = _ensure_style(doc, name)
    _set_style_font(
        style,
        font_name,
        size,
        bold=bold,
        italic=italic,
        color=color,
    )
    paragraph = style.paragraph_format
    paragraph.alignment = alignment
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    paragraph.keep_with_next = keep_with_next
    paragraph.keep_together = keep_together
    paragraph.first_line_indent = Pt(first_line_indent)
    paragraph.left_indent = Pt(left_indent)
    paragraph.widow_control = True
    _set_outline_level(style, outline_level)
    return style


def _configure_academic_document(
    doc: Document,
    layout_profile: str = _DEFAULT_LAYOUT_PROFILE,
) -> None:
    if layout_profile not in _LAYOUT_PROFILES:
        raise ValueError(
            f"unknown layout profile {layout_profile!r}; "
            f"choose from {', '.join(_LAYOUT_PROFILES)}"
        )
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    if layout_profile == "chemvellum_journal":
        section.top_margin = Inches(_CHEMVELLUM["top_margin"])
        section.bottom_margin = Inches(_CHEMVELLUM["bottom_margin"])
        section.left_margin = Inches(_CHEMVELLUM["side_margin"])
        section.right_margin = Inches(_CHEMVELLUM["side_margin"])
        section.header_distance = Inches(_CHEMVELLUM["header_distance"])
        section.footer_distance = Inches(_CHEMVELLUM["footer_distance"])
    else:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)

    if layout_profile == "legacy_report":
        _configure_style(
            doc, _S["title"], size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            before=0, after=12, line_spacing=1.15,
        )
        _configure_style(doc, _S["author"], size=11.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=3)
        _configure_style(doc, _S["address"], size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=3)
        _configure_style(doc, _S["abstract"], size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, line_spacing=1.15)
        _configure_style(doc, _S["keywords"], size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=8, line_spacing=1.15)
        _configure_style(doc, _S["body"], size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line_spacing=1.5)
        _configure_style(doc, _S["body_lead"], size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line_spacing=1.5)
        _configure_style(doc, _S["list"], size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, line_spacing=1.5)
        _configure_style(doc, _S["toc"], size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=3, line_spacing=1.15)
        _configure_style(doc, _S["h1"], size=14, bold=True, before=12, after=6, keep_with_next=True, outline_level=0)
        _configure_style(doc, _S["h2"], size=12, bold=True, before=10, after=4, keep_with_next=True, outline_level=1)
        _configure_style(doc, _S["h3"], size=11, bold=True, italic=True, before=8, after=3, keep_with_next=True, outline_level=2)
        _configure_style(doc, _S["figure"], size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=8, keep_together=True)
        _configure_style(doc, _S["figure_display"], size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2, keep_with_next=True, keep_together=True)
        _configure_style(doc, _S["table_title"], size=10, bold=True, before=6, after=4, keep_with_next=True, keep_together=True)
        _configure_style(doc, _S["table_header"], size=9.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _configure_style(doc, _S["table_body"], size=9.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _configure_style(
            doc, _S["references"], size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            after=3, left_indent=36, first_line_indent=-36,
        )
        normal_font, normal_size, normal_after, normal_spacing = "Times New Roman", 12, 6, 1.5
    elif layout_profile == "chemvellum_journal":
        # ChemVellum Journal is a Word-native editorial system: a compact
        # two-column reading grid, restrained teal identity, and full-width
        # evidence figures/tables.  It intentionally stays distinct from any
        # one publisher's masthead or page furniture.
        _configure_style(
            doc, _S["title"], size=21.5, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8,
            line_spacing=1.0, keep_together=True,
            font_name="Georgia", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["author"], size=8.8, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            after=2, font_name="Arial", color=_CHEMVELLUM["muted"],
        )
        _configure_style(
            doc, _S["address"], size=8.2, italic=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, after=2,
            font_name="Arial", color=_CHEMVELLUM["muted"],
        )
        _configure_style(
            doc, _S["abstract"], size=10.0,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=6, after=9,
            line_spacing=1.10, left_indent=13,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["keywords"], size=8.2,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, after=7,
            line_spacing=1.0, font_name="Arial", color=_CHEMVELLUM["muted"],
        )
        _configure_style(
            doc, _S["body"], size=9.5,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=1.5,
            line_spacing=1.04, first_line_indent=9,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["body_lead"], size=9.5,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=1.5,
            line_spacing=1.04, first_line_indent=9,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["list"], size=9.0,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, after=1.5,
            line_spacing=1.02, font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["toc"], size=7.8,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, after=0.5,
            line_spacing=1.0, font_name="Arial", color=_CHEMVELLUM["muted"],
        )
        _configure_style(
            doc, _S["h1"], size=11.5, bold=True, before=8, after=3,
            line_spacing=1.0, keep_with_next=True, outline_level=0,
            font_name="Arial", color=_CHEMVELLUM["teal"],
        )
        _configure_style(
            doc, _S["h2"], size=9.5, bold=True, before=6, after=2.2,
            line_spacing=1.0, keep_with_next=True, outline_level=1,
            font_name="Arial", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["h3"], size=9.0, bold=True, italic=True,
            before=5, after=1.5, line_spacing=1.0,
            keep_with_next=True, outline_level=2,
            font_name="Arial", color=_CHEMVELLUM["muted"],
        )
        _configure_style(
            doc, _S["figure"], size=8.0,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6,
            line_spacing=1.04, keep_together=True,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["figure_display"], size=9.5,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=2,
            keep_with_next=True, keep_together=True,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["table_title"], size=8.0,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=2,
            keep_with_next=True, keep_together=True,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["table_header"], size=7.5, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0,
            font_name="Arial", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["table_body"], size=7.5,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        _configure_style(
            doc, _S["references"], size=7.5,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, after=1,
            line_spacing=1.0, left_indent=18, first_line_indent=-18,
            font_name="Cambria", color=_CHEMVELLUM["ink"],
        )
        normal_font, normal_size, normal_after, normal_spacing = "Cambria", 9.5, 1.5, 1.04
    else:
        # A restrained journal-like house style.  Styles carry typography;
        # direct formatting is reserved for semantic emphasis and scripts.
        _configure_style(
            doc, _S["title"], size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            before=0, after=10, line_spacing=1.08, keep_together=True,
        )
        _configure_style(doc, _S["author"], size=10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        _configure_style(doc, _S["address"], size=9.5, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        _configure_style(doc, _S["abstract"], size=10.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=5, line_spacing=1.08)
        _configure_style(doc, _S["keywords"], size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=7, line_spacing=1.08)
        _configure_style(
            doc, _S["body"], size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            after=0, line_spacing=1.15, first_line_indent=11,
        )
        _configure_style(
            doc, _S["body_lead"], size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            after=0, line_spacing=1.15,
        )
        _configure_style(
            doc, _S["list"], size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            after=2, line_spacing=1.08,
        )
        _configure_style(doc, _S["toc"], size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=1, line_spacing=1.0)
        _configure_style(doc, _S["h1"], size=13.5, bold=True, before=11, after=4, keep_with_next=True, outline_level=0)
        _configure_style(doc, _S["h2"], size=11.5, bold=True, before=9, after=3, keep_with_next=True, outline_level=1)
        _configure_style(doc, _S["h3"], size=10.5, bold=True, italic=True, before=7, after=2, keep_with_next=True, outline_level=2)
        _configure_style(doc, _S["figure"], size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=6, keep_together=True)
        _configure_style(doc, _S["figure_display"], size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2, keep_with_next=True, keep_together=True)
        _configure_style(doc, _S["table_title"], size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3, keep_with_next=True, keep_together=True)
        _configure_style(doc, _S["table_header"], size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
        _configure_style(doc, _S["table_body"], size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
        _configure_style(
            doc, _S["references"], size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            after=2, line_spacing=1.0, left_indent=36, first_line_indent=-36,
        )
        normal_font, normal_size, normal_after, normal_spacing = "Times New Roman", 11, 0, 1.15

    normal = doc.styles["Normal"]
    _set_style_font(normal, normal_font, normal_size)
    normal.paragraph_format.space_after = Pt(normal_after)
    normal.paragraph_format.line_spacing = normal_spacing
    if layout_profile == "chemvellum_journal":
        _set_section_columns(section, 1, _CHEMVELLUM["column_gap_dxa"])


def _append_page_number(
    paragraph,
    *,
    align_right: bool = True,
    font_name: str = "Times New Roman",
    size_pt: float = 9,
) -> None:
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    _set_word_run_font(run, font_name, size_pt)


def _set_run_width_scale(run, percent: int) -> None:
    rpr = run._element.get_or_add_rPr()
    width = rpr.find(qn("w:w"))
    if width is None:
        width = OxmlElement("w:w")
        rpr.append(width)
    width.set(qn("w:val"), str(percent))


def _set_paragraph_border(
    paragraph,
    edge: str,
    *,
    color: str,
    size: int = 4,
    space: int = 3,
) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    for existing in list(borders.findall(qn(f"w:{edge}"))):
        borders.remove(existing)
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for existing in list(ppr.findall(qn("w:shd"))):
        ppr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    ppr.append(shading)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _configure_chemvellum_header_footer(doc: Document) -> None:
    """Apply the original ChemVellum wordmark and quiet running furniture."""
    first = doc.sections[0]
    for section in doc.sections[1:]:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    header_p = first.header.paragraphs[0]
    _clear_paragraph(header_p)
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.line_spacing = 1.0
    header_p.paragraph_format.tab_stops.add_tab_stop(
        Inches(_CHEMVELLUM["content_width"]), WD_TAB_ALIGNMENT.RIGHT
    )

    chem = header_p.add_run("CHEM")
    _set_word_run_font(chem, "Georgia", 12.2)
    chem.bold = True
    chem.font.color.rgb = RGBColor.from_string(_CHEMVELLUM["teal"])
    _set_run_width_scale(chem, 82)
    vellum = header_p.add_run("VELLUM")
    _set_word_run_font(vellum, "Georgia", 12.2)
    vellum.bold = True
    vellum.font.color.rgb = RGBColor.from_string(_CHEMVELLUM["ink"])
    _set_run_width_scale(vellum, 82)
    label = header_p.add_run("\tR E V I E W")
    _set_word_run_font(label, "Arial", 7.2)
    label.bold = True
    label.font.color.rgb = RGBColor.from_string(_CHEMVELLUM["teal"])

    footer_p = first.footer.paragraphs[0]
    _clear_paragraph(footer_p)
    footer_p.paragraph_format.space_before = Pt(3)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.paragraph_format.line_spacing = 1.0
    footer_p.paragraph_format.tab_stops.add_tab_stop(
        Inches(_CHEMVELLUM["content_width"]), WD_TAB_ALIGNMENT.RIGHT
    )
    _set_paragraph_border(
        footer_p,
        "top",
        color=_CHEMVELLUM["rule"],
        size=4,
        space=3,
    )
    descriptor = footer_p.add_run(
        "CHEMVELLUM  /  EVIDENCE-LED CHEMICAL REVIEW SYNTHESIS"
    )
    _set_word_run_font(descriptor, "Arial", 6.4)
    descriptor.font.color.rgb = RGBColor.from_string(_CHEMVELLUM["muted"])
    footer_p.add_run("\t")
    _append_page_number(
        footer_p,
        align_right=False,
        font_name="Arial",
        size_pt=7.0,
    )


def _configure_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        _append_page_number(paragraph)


def _set_section_columns(section, count: int, space_dxa: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    for column in list(cols.findall(qn("w:col"))):
        cols.remove(column)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_dxa))
    cols.set(qn("w:equalWidth"), "1")


def _add_chemvellum_section(
    doc: Document,
    columns: int,
    start_type=WD_SECTION.CONTINUOUS,
):
    section = doc.add_section(start_type)
    section.page_width = Inches(_CHEMVELLUM["page_width"])
    section.page_height = Inches(_CHEMVELLUM["page_height"])
    section.top_margin = Inches(_CHEMVELLUM["top_margin"])
    section.bottom_margin = Inches(_CHEMVELLUM["bottom_margin"])
    section.left_margin = Inches(_CHEMVELLUM["side_margin"])
    section.right_margin = Inches(_CHEMVELLUM["side_margin"])
    section.header_distance = Inches(_CHEMVELLUM["header_distance"])
    section.footer_distance = Inches(_CHEMVELLUM["footer_distance"])
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    _set_section_columns(section, columns, _CHEMVELLUM["column_gap_dxa"])
    return section


def _writable_path(path: Path) -> str:
    """Use Windows extended-length syntax without changing the visible filename."""
    resolved = str(path.resolve())
    if os.name != "nt" or resolved.startswith("\\\\?\\"):
        return resolved
    if resolved.startswith("\\\\"):
        return "\\\\?\\UNC\\" + resolved[2:]
    return "\\\\?\\" + resolved

_UNICODE_SUPERSCRIPT_MAP: Dict[str, str] = {
    "⁰": "0",
    "¹": "1",
    "²": "2",
    "³": "3",
    "⁴": "4",
    "⁵": "5",
    "⁶": "6",
    "⁷": "7",
    "⁸": "8",
    "⁹": "9",
    "⁺": "+",
    "⁻": "-",
    "⁼": "=",
    "⁽": "(",
    "⁾": ")",
}

_UNICODE_SUBSCRIPT_MAP: Dict[str, str] = {
    "₀": "0",
    "₁": "1",
    "₂": "2",
    "₃": "3",
    "₄": "4",
    "₅": "5",
    "₆": "6",
    "₇": "7",
    "₈": "8",
    "₉": "9",
    "₊": "+",
    "₋": "-",
    "₌": "=",
    "₍": "(",
    "₎": ")",
}

# ---------------------------------------------------------------------------
# Run dataclass
# ---------------------------------------------------------------------------

@dataclass
class Run:
    text:        str  = ""
    bold:        bool = False
    italic:      bool = False
    code:        bool = False
    superscript: bool = False
    subscript:   bool = False
    math:        str  = ""

# ---------------------------------------------------------------------------
# Inline parser
# ---------------------------------------------------------------------------

_STACKED_SCRIPT_TOKEN_RE = re.compile(
    r"([A-Za-z][A-Za-z0-9]*)_([^_\s\n]+)_\^([^\^\s\n]+)\^"
)

_INLINE_RE = re.compile(
    r"(\$\$[\s\S]*?\$\$"
    r"|\$[^$\n]+?\$"
    r"|\*\*\*(?:\S[^*\n]*?\S|\S)\*\*\*"
    r"|\*\*(?:\S[^*\n]*?\S|\S)\*\*"
    r"|[A-Za-z][A-Za-z0-9]*_[^_\s\n]+_\^[^\^\s\n]+\^"
    r"|\^[^\^\s\n]+?\^"
    r"|_[^_\s\n]+_"
    r"|\*(?:\S[^*\n]*?\S|\S)\*"
    r"|__(?:\S[^_\n]*?\S|\S)__"
    r"|_(?:\S[^_\n]*?\S|\S)_"
    r"|`[^`\n]+`"
    r"|\[@([^\]]+)\]"
    r"|\[([^\]]*)\]\([^\)]*\)"
    r")"
)


def _parse_nested(inner: str, bold: bool = False, italic: bool = False) -> List[Run]:
    runs = parse_inline(inner)
    for r in runs:
        if bold:
            r.bold = True
        if italic:
            r.italic = True
    return runs


def parse_inline(raw: str) -> List[Run]:
    # Review drafts commonly express isotopes as ``^15_N``.  Normalize that
    # compact chemistry notation to the explicit superscript form understood
    # by the inline parser while leaving the element/formula text untouched.
    raw = re.sub(r"\^(\d+)_([A-Z])", r"^\1^\2", raw)
    runs: List[Run] = []
    pos = 0
    for m in _INLINE_RE.finditer(raw):
        if m.start() > pos:
            runs.append(Run(text=raw[pos:m.start()]))
        token = m.group(0)
        char_before = raw[m.start() - 1] if m.start() > 0 else ""
        stacked_script = _STACKED_SCRIPT_TOKEN_RE.fullmatch(token)

        if token.startswith("$$"):
            runs.append(Run(math=token[2:-2].strip()))
        elif token.startswith("$"):
            runs.append(Run(math=token[1:-1].strip()))
        elif stacked_script:
            base, subscript, superscript = stacked_script.groups()
            if superscript.startswith("-"):
                superscript = "−" + superscript[1:]
            runs.append(
                Run(math=f"{base}_{{{subscript}}}^{{{superscript}}}")
            )
        elif token.startswith("***"):
            runs.extend(_parse_nested(token[3:-3], bold=True, italic=True))
        elif token.startswith("**"):
            runs.extend(_parse_nested(token[2:-2], bold=True))
        elif token.startswith("^") and token.endswith("^"):
            runs.append(Run(text=token[1:-1], superscript=True))
        elif token.startswith("_") and token.endswith("_") and " " not in token[1:-1]:
            if char_before.isalnum() or char_before in ")]}":
                runs.append(Run(text=token[1:-1], subscript=True))
            else:
                runs.append(Run(text=token[1:-1], italic=True))
        elif token.startswith("*"):
            runs.extend(_parse_nested(token[1:-1], italic=True))
        elif token.startswith("__"):
            runs.extend(_parse_nested(token[2:-2], bold=True))
        elif token.startswith("_"):
            runs.extend(_parse_nested(token[1:-1], italic=True))
        elif token.startswith("`"):
            runs.append(Run(text=token[1:-1], code=True))
        elif token.startswith("[@"):
            cite_key = m.group(1) or token[2:-1]
            runs.append(Run(text=f"[{cite_key}]"))
        elif token.startswith("["):
            display = m.group(2)
            runs.append(Run(text=display if display is not None else token))
        else:
            runs.append(Run(text=token))
        pos = m.end()

    if pos < len(raw):
        runs.append(Run(text=raw[pos:]))
    return runs or [Run(text=raw)]


# ---------------------------------------------------------------------------
# Font + run application
# ---------------------------------------------------------------------------

def _math_run(text: str, font_name: str, size_pt: float):
    run = OxmlElement("m:r")
    math_properties = OxmlElement("m:rPr")
    math_style = OxmlElement("m:sty")
    math_style.set(qn("m:val"), "p")
    math_properties.append(math_style)
    run.append(math_properties)

    word_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), font_name)
    word_properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(round(size_pt * 2))))
    word_properties.append(size)
    run.append(word_properties)

    text_element = OxmlElement("m:t")
    text_element.text = text
    run.append(text_element)
    return run


def _append_stacked_script(
    para,
    base: str,
    subscript: str,
    superscript: str,
    font_name: str,
    size_pt: float,
) -> None:
    """Append one compact Word math object with paired sub- and superscripts."""
    math = OxmlElement("m:oMath")
    stacked = OxmlElement("m:sSubSup")
    for tag, text in (
        ("m:e", base),
        ("m:sub", subscript),
        ("m:sup", superscript),
    ):
        container = OxmlElement(tag)
        container.append(_math_run(text, font_name, size_pt))
        stacked.append(container)
    math.append(stacked)
    para._p.append(math)


def _apply_math(
    para,
    latex: str,
    font_name: str = "Times New Roman",
    size_pt: float = 12,
) -> None:
    if _LATEX_OK:
        try:
            LatexToWordElement(latex).add_latex_to_paragraph(para)
            return
        except Exception:
            pass
    expression = latex.strip()
    chemical = re.fullmatch(r"\\ce\{(.+)\}", expression)
    if chemical:
        expression = chemical.group(1)
    expression = re.sub(r"\\(?:mathrm|text)\{([^{}]*)\}", r"\1", expression)
    stacked_script = re.fullmatch(
        r"([A-Za-z][A-Za-z0-9]*)_\{([^{}]+)\}\^\{([^{}]+)\}",
        expression,
    )
    if stacked_script:
        _append_stacked_script(
            para,
            *stacked_script.groups(),
            font_name,
            size_pt,
        )
        return
    pattern = re.compile(r"([_^])(?:\{([^{}]+)\}|([A-Za-z0-9+\-=]+))")
    position = 0
    for match in pattern.finditer(expression):
        if match.start() > position:
            run = para.add_run(expression[position:match.start()])
            _set_word_run_font(run, font_name, size_pt)
        run = para.add_run(match.group(2) or match.group(3) or "")
        _set_word_run_font(run, font_name, size_pt)
        run.font.subscript = match.group(1) == "_"
        run.font.superscript = match.group(1) == "^"
        position = match.end()
    if position < len(expression):
        run = para.add_run(expression[position:])
        _set_word_run_font(run, font_name, size_pt)


def _split_script_segments(text: str) -> List[Tuple[str, str]]:
    segments: List[Tuple[str, str]] = []
    mode = "normal"
    current = ""

    def looks_like_ascii_subscript(index: int) -> bool:
        if index < 0 or index >= len(text):
            return False
        char = text[index]
        if not char.isdigit() or index == 0:
            return False
        prev = text[index - 1]
        if prev in "-–—/[ ":
            return False
        if prev.isalpha():
            return True
        if prev in ")]}" and index >= 2 and text[index - 2].isalpha():
            return True
        return False

    def flush() -> None:
        nonlocal current
        if current:
            segments.append((mode, current))
            current = ""

    for idx, char in enumerate(text):
        if char in _UNICODE_SUPERSCRIPT_MAP:
            char_mode = "superscript"
            rendered = _UNICODE_SUPERSCRIPT_MAP[char]
        elif char in _UNICODE_SUBSCRIPT_MAP:
            char_mode = "subscript"
            rendered = _UNICODE_SUBSCRIPT_MAP[char]
        elif looks_like_ascii_subscript(idx):
            char_mode = "subscript"
            rendered = char
        else:
            char_mode = "normal"
            rendered = char
        if char_mode != mode:
            flush()
            mode = char_mode
        current += rendered
    flush()
    return segments or [("normal", text)]


def _split_script_segments_v2(text: str) -> List[Tuple[str, str]]:
    """Apply explicit Unicode scripts and conservative chemistry-aware defaults.

    Explicit Markdown `_sub_` and `^super^` is handled before this function.
    Automatic conversion is intentionally limited to multi-element formulae,
    sp2/sp3 hybridization, SN1/SN2 notation, and hapticity labels such as
    eta1/eta3 (written with the Greek eta character) so years and ordinary
    alphanumeric labels stay intact.
    """
    segments: List[Tuple[str, str]] = []

    def append(mode: str, value: str) -> None:
        if not value:
            return
        if segments and segments[-1][0] == mode:
            segments[-1] = (mode, segments[-1][1] + value)
        else:
            segments.append((mode, value))

    def append_unicode(value: str) -> None:
        for char in value:
            if char in _UNICODE_SUPERSCRIPT_MAP:
                append("superscript", _UNICODE_SUPERSCRIPT_MAP[char])
            elif char in _UNICODE_SUBSCRIPT_MAP:
                append("subscript", _UNICODE_SUBSCRIPT_MAP[char])
            else:
                append("normal", char)

    parenthesized_complex = (
        r"(?:[A-Z][a-z]?\d*)+(?:\([A-Za-z][A-Za-z0-9]*\)\d*)+"
        r"(?:[A-Z][a-z]?\d*)*"
    )
    token_re = re.compile(
        rf"(?<![A-Za-z0-9])(?:η[1-9]|sp[23]|SN[12]|{parenthesized_complex}|"
        rf"(?:[A-Z][a-z]?\d*){{2,}})(?![A-Za-z0-9])"
    )
    position = 0
    for match in token_re.finditer(text):
        append_unicode(text[position:match.start()])
        token = match.group(0)
        if re.fullmatch(r"η[1-9]", token):
            append("normal", "η")
            append("superscript", token[-1])
        elif re.fullmatch(r"sp[23]", token):
            append("normal", "sp")
            append("superscript", token[-1])
        elif re.fullmatch(r"SN[12]", token):
            append("normal", "S")
            append("subscript", token[1:])
        elif re.fullmatch(parenthesized_complex, token):
            for piece in re.finditer(r"\d+|[^\d]+", token):
                append("subscript" if piece.group(0).isdigit() else "normal", piece.group(0))
        else:
            for piece in re.finditer(r"[A-Z][a-z]?|\d+", token):
                append("subscript" if piece.group(0).isdigit() else "normal", piece.group(0))
        position = match.end()
    append_unicode(text[position:])
    return segments or [("normal", text)]


def _set_word_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def apply_runs(
    para,
    runs: List[Run],
    spec_key: str = "body",
    force_bold: bool = False,
    force_italic: bool = False,
) -> None:
    spec        = _FONT_SPEC.get(spec_key, _FONT_SPEC["body"])
    style_font  = para.style.font
    font_name   = style_font.name or spec["font"]
    size_pt     = style_font.size.pt if style_font.size else spec["size"]
    for r in runs:
        if r.math:
            _apply_math(para, r.math, font_name, size_pt)
            continue
        segments = [("normal", r.text)] if r.code else _split_script_segments_v2(r.text)
        for segment_mode, segment_text in segments:
            if not segment_text:
                continue
            wr = para.add_run(segment_text)
            if r.code:
                _set_word_run_font(wr, "Courier New", 9)
            # Paragraph styles own the base font, size, weight, and alignment.
            # Only meaning-bearing inline emphasis remains direct formatting.
            if force_bold or r.bold:
                wr.bold = True
            if force_italic or r.italic:
                wr.italic = True
            if r.superscript or segment_mode == "superscript":
                wr.font.superscript = True
            if r.subscript or segment_mode == "subscript":
                wr.font.subscript = True


_CAPTION_LABEL_RE = re.compile(
    r"^((?:Figure|Fig\.|Table|Scheme|Chart)\s+\d+[A-Za-z]?(?:[.:])?)\s*(.*)$",
    re.I,
)


def _apply_caption_runs(paragraph, inline_text: str, spec_key: str) -> None:
    """Bold only the semantic caption label, not the full caption sentence."""
    match = _CAPTION_LABEL_RE.match(inline_text.strip())
    if not match:
        apply_runs(paragraph, parse_inline(inline_text), spec_key=spec_key)
        return
    label, remainder = match.groups()
    first_label_run = len(paragraph.runs)
    apply_runs(paragraph, parse_inline(label), spec_key=spec_key, force_bold=True)
    for run in paragraph.runs[first_label_run:]:
        run.font.color.rgb = RGBColor.from_string(_CHEMVELLUM["teal"])
    if remainder:
        paragraph.add_run(" ")
        apply_runs(paragraph, parse_inline(remainder), spec_key=spec_key)


# ---------------------------------------------------------------------------
# Paragraph factory
# ---------------------------------------------------------------------------

def _para(
    doc: Document,
    style_key: str,
    spec_key: str,
    inline_text: str = "",
    force_bold: bool = False,
    force_italic: bool = False,
):
    p = doc.add_paragraph(style=_S.get(style_key, _S["body"]))
    if inline_text:
        if style_key in {"figure", "scheme", "chart", "table_title"} and not (force_bold or force_italic):
            _apply_caption_runs(p, inline_text, spec_key)
        else:
            apply_runs(p, parse_inline(inline_text),
                       spec_key=spec_key,
                       force_bold=force_bold,
                       force_italic=force_italic)
    return p


def _reference_para(
    doc: Document,
    number: int,
    inline_text: str,
    layout_profile: str = _DEFAULT_LAYOUT_PROFILE,
):
    """Add a stable, hanging-indent bibliography entry.

    Citation merging has already fixed the visible reference numbers in the
    Markdown source.  Keeping those labels as text avoids Word's tendency to
    continue a bibliography from the nearest visually similar body list.
    """
    p = _para(doc, "references", "references", f"\t{number}.\t{inline_text}")
    if layout_profile == "chemvellum_journal":
        right_tab, left_tab = 0.16, 0.21
    else:
        right_tab, left_tab = 0.42, 0.50
    p.paragraph_format.tab_stops.add_tab_stop(Inches(right_tab), WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(left_tab), WD_TAB_ALIGNMENT.LEFT)
    return p


def _next_numbering_id(numbering, tag: str, attribute: str) -> int:
    values = []
    for element in numbering.findall(qn(tag)):
        raw = element.get(qn(attribute))
        if raw and raw.isdigit():
            values.append(int(raw))
    return max(values, default=0) + 1


def _create_numbering_definition(
    doc: Document,
    ordered: bool,
    reference: bool = False,
    start_at: int = 1,
) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = _next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{(0x6A1D0000 + abstract_id) & 0xFFFFFFFF:08X}")
    abstract.append(nsid)
    # Word also uses the template code when deciding whether separately
    # defined lists should continue one another.  Give every definition a
    # distinct template identity so a conclusion list cannot flow into the
    # bibliography merely because their level formatting matches.
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"{(0x7B2E0000 + abstract_id) & 0xFFFFFFFF:08X}")
    abstract.append(template)
    multi = OxmlElement("w:multiLevelType")
    # References are deliberately single-level.  Word otherwise recognizes
    # two visually similar multilevel definitions as one continued sequence
    # when it opens the file, even when their numId, nsid, template, and
    # startOverride values differ.
    multi.set(qn("w:val"), "singleLevel" if reference else "multilevel")
    abstract.append(multi)
    level_count = 1 if reference else 3
    for level in range(level_count):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), str(max(1, start_at)))
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), f"%{level + 1}." if ordered else "•")
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        left = (720 if reference else 540) + level * 360
        hanging = 360 if reference else 270
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), str(left))
        indentation.set(qn("w:hanging"), str(hanging))
        ppr.extend([tabs, indentation])
        lvl.extend([start, num_fmt, lvl_text, suffix, justification, ppr])
        abstract.append(lvl)
    numbering.append(abstract)

    num_id = _next_numbering_id(numbering, "w:num", "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if ordered and not reference:
        for level in range(level_count):
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), str(level))
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), str(max(1, start_at)))
            override.append(start_override)
            num.append(override)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, min(level, 2))))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------

def _set_cell_borders(
    cell,
    *,
    top: bool = False,
    bottom: bool = False,
    color: str = "000000",
) -> None:
    """Apply a white three-line-table border treatment."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        for existing in list(borders.findall(qn(f"w:{edge}"))):
            borders.remove(existing)
        elem = OxmlElement(f"w:{edge}")
        enabled = (edge == "top" and top) or (edge == "bottom" and bottom)
        elem.set(qn("w:val"), "single" if enabled else "nil")
        elem.set(qn("w:sz"), "8" if enabled else "0")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)
        borders.append(elem)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_row_pagination(row, *, repeat_header: bool = False) -> None:
    """Keep a row intact and optionally repeat it as the table header."""
    tr_pr = row._tr.get_or_add_trPr()
    for tag in ("w:cantSplit", "w:tblHeader"):
        for existing in list(tr_pr.findall(qn(tag))):
            tr_pr.remove(existing)
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)
    if repeat_header:
        table_header = OxmlElement("w:tblHeader")
        table_header.set(qn("w:val"), "true")
        tr_pr.append(table_header)


def _set_row_keep_with_next(row, value: bool) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = value


def _protect_table_orphans(table) -> None:
    """Keep captions/headers with data and avoid one-row table tails."""
    if not table.rows:
        return
    body_count = len(table.rows) - 1
    _set_row_keep_with_next(table.rows[0], body_count > 0)
    if body_count >= 2:
        _set_row_keep_with_next(table.rows[1], True)
        _set_row_keep_with_next(table.rows[-2], True)


def _table_header_key(raw: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", raw.lower()).strip()


def _chemvellum_table_widths(
    header: List[str],
    ncols: int,
    width_inches: float,
) -> Optional[List[float]]:
    """Return deliberate widths for common chemistry comparison columns."""
    semantic_widths = []
    for raw in (header + [""] * ncols)[:ncols]:
        key = _table_header_key(raw)
        if "catalyst" in key or "material" in key:
            semantic_widths.append(1.30)
        elif "electrolyte" in key or "medium" in key:
            semantic_widths.append(1.85)
        elif "operating" in key or "potential" in key or "current density" in key:
            semantic_widths.append(1.05)
        elif "faradaic" in key or re.search(r"(?:^| )fe(?: |$)", key):
            semantic_widths.append(0.70)
        elif "yield" in key or "rate" in key:
            semantic_widths.append(1.30)
        elif "stability" in key or "duration" in key:
            semantic_widths.append(0.62)
        elif key.startswith("ref") or "reference" in key:
            semantic_widths.append(0.38)
        else:
            return None
    scale = width_inches / sum(semantic_widths)
    return [width * scale for width in semantic_widths]


def _table_column_alignment(header_text: str):
    key = _table_header_key(header_text)
    centered_tokens = (
        "operating",
        "potential",
        "current density",
        "faradaic",
        "yield",
        "rate",
        "stability",
        "duration",
        "ref",
    )
    if any(token in key for token in centered_tokens) or re.search(r"(?:^| )fe(?: |$)", key):
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT


def _add_table_single(
    doc: Document,
    header: List[str],
    rows: List[List[str]],
    *,
    width_inches: float = 6.5,
    layout_profile: str = _DEFAULT_LAYOUT_PROFILE,
) -> None:
    ncols = max(len(header), max((len(r) for r in rows), default=1))
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_row_pagination(table.rows[0], repeat_header=True)
    for row in table.rows[1:]:
        _set_table_row_pagination(row)
    content_rows = [header] + rows
    widths = None
    if layout_profile == "chemvellum_journal":
        widths = _chemvellum_table_widths(header, ncols, width_inches)
    if widths is None:
        weights = []
        for index in range(ncols):
            longest = max((len(str(row[index])) if index < len(row) else 0 for row in content_rows), default=1)
            weights.append(max(8, min(longest, 50)))
        total_weight = sum(weights) or ncols
        minimum_column_width = 0.38 if layout_profile == "chemvellum_journal" else 0.8
        widths = [max(minimum_column_width, width_inches * weight / total_weight) for weight in weights]
        scale = width_inches / sum(widths)
        widths = [width * scale for width in widths]
    dxa_widths = [int(round(width * 1440)) for width in widths]
    table_width_dxa = int(round(width_inches * 1440))
    dxa_widths[-1] += table_width_dxa - sum(dxa_widths)

    table_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for existing in list(table_pr.findall(qn(tag))):
            table_pr.remove(existing)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(table_width_dxa))
    table_width.set(qn("w:type"), "dxa")
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_pr.extend([table_width, table_indent, layout])
    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for index, grid_column in enumerate(grid_columns[:ncols]):
        grid_column.set(qn("w:w"), str(dxa_widths[index]))
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.width = Inches(widths[j])
        cell._tc.get_or_add_tcPr().get_or_add_tcW().set(qn("w:w"), str(dxa_widths[j]))
        cell._tc.get_or_add_tcPr().get_or_add_tcW().set(qn("w:type"), "dxa")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        cell.paragraphs[0].style = doc.styles[_S["table_header"]]
        cell.paragraphs[0].alignment = _table_column_alignment(h)
        apply_runs(cell.paragraphs[0], parse_inline(h),
                   spec_key="table_header", force_bold=True)
        border_color = _CHEMVELLUM["teal"] if layout_profile == "chemvellum_journal" else "000000"
        _set_cell_borders(cell, top=True, bottom=True, color=border_color)
        if layout_profile == "chemvellum_journal":
            _set_cell_shading(cell, _CHEMVELLUM["pale"])
            _set_cell_margins(cell, top=60, start=80, bottom=60, end=80)
        else:
            _set_cell_margins(cell)
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i + 1, j)
            cell.width = Inches(widths[j])
            cell._tc.get_or_add_tcPr().get_or_add_tcW().set(qn("w:w"), str(dxa_widths[j]))
            cell._tc.get_or_add_tcPr().get_or_add_tcW().set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            cell.paragraphs[0].style = doc.styles[_S["table_body"]]
            cell.paragraphs[0].alignment = _table_column_alignment(
                header[j] if j < len(header) else ""
            )
            apply_runs(cell.paragraphs[0],
                       parse_inline(row[j] if j < len(row) else ""),
                       spec_key="table_body")
            border_color = _CHEMVELLUM["teal"] if layout_profile == "chemvellum_journal" else "000000"
            _set_cell_borders(
                cell,
                bottom=i == len(rows) - 1,
                color=border_color,
            )
            if layout_profile == "chemvellum_journal":
                if i % 2 == 1:
                    _set_cell_shading(cell, _CHEMVELLUM["row_pale"])
                _set_cell_margins(cell, top=35, start=80, bottom=35, end=80)
            else:
                _set_cell_margins(cell)
    _protect_table_orphans(table)


def _add_table(
    doc: Document,
    header: List[str],
    rows: List[List[str]],
    *,
    width_inches: float = 6.5,
    layout_profile: str = _DEFAULT_LAYOUT_PROFILE,
) -> None:
    """Add a readable table, splitting overly wide comparison grids.

    A method table with six or more narrow columns is technically inside
    the page but practically unreadable.  Repeat the identifying first column
    and split the remaining fields into compact continuation tables instead of
    shrinking words into vertical fragments.
    """
    ncols = max(len(header), max((len(row) for row in rows), default=1))
    maximum_columns = 7 if layout_profile == "chemvellum_journal" else 5
    if ncols <= maximum_columns:
        _add_table_single(
            doc,
            header,
            rows,
            width_inches=width_inches,
            layout_profile=layout_profile,
        )
        return
    padded_header = header + [""] * (ncols - len(header))
    padded_rows = [row + [""] * (ncols - len(row)) for row in rows]
    for chunk_index, start in enumerate(range(1, ncols, 4)):
        indices = [0] + list(range(start, min(start + 4, ncols)))
        if chunk_index:
            _para(doc, "table_title", "table_title", "Table continued", force_bold=True)
        _add_table_single(
            doc,
            [padded_header[index] for index in indices],
            [[row[index] for index in indices] for row in padded_rows],
            width_inches=width_inches,
            layout_profile=layout_profile,
        )


# ---------------------------------------------------------------------------
# Block tokenizer
# ---------------------------------------------------------------------------

@dataclass
class Block:
    kind:     str
    level:    int             = 0
    text:     str             = ""
    ordered:  bool            = False
    list_number: int          = 0
    depth:    int             = 0
    code:     str             = ""
    language: str             = ""
    header:   List[str]       = field(default_factory=list)
    rows:     List[List[str]] = field(default_factory=list)
    alt:      str             = ""
    path:     str             = ""
    latex:    str             = ""
    lines:    List[str]       = field(default_factory=list)


_HEADING_RE    = re.compile(r"^(#{1,6})\s+(.*)")
_EMBEDDED_HEADING_PREFIX_RE = re.compile(r"^#{1,6}\s+")
_NUMBERED_SECTION_HEADING_RE = re.compile(r"^\d+(?:\.\d+)*\.\s+\S")
_HTML_ANCHOR_RE = re.compile(r"^<a\s+id=[\"']ref-\d+[\"']\s*>\s*</a>\s*$", re.I)
_HTML_COMMENT_START_RE = re.compile(r"^\s*<!--")
_UL_RE         = re.compile(r"^(\s*)[-*]\s+(.*)")
_OL_RE         = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)")
_FENCE_RE      = re.compile(r"^```(\w*)\s*$")
_MATH_FENCE_RE = re.compile(r"^\$\$\s*$")
_IMG_RE        = re.compile(r"^!\[((?:\\.|[^\]])*)\]\(([^)]+)\)\s*$")
_TABLE_ROW_RE  = re.compile(r"^\|.+")
_HR_RE         = re.compile(r"^(?:-{3,}|_{3,}|\*{3,})\s*$")
_REF_ENTRY_RE  = re.compile(r"^(?:\d+[.)]\s+|\[\d+\]\s+\S|\[@[^\]]+\]:)")
_AFFIL_START   = re.compile(r"^\^[0-9,]+\^")
_INDENTED_RE   = re.compile(r"^(?: {4,}|\t+)(.*)$")


def _is_continuation(line: str) -> bool:
    if not line.strip():
        return False
    if _REF_ENTRY_RE.match(line):
        return False
    if _AFFIL_START.match(line):
        return False
    if _HTML_COMMENT_START_RE.match(line):
        return False
    for pat in (_HEADING_RE, _FENCE_RE, _TABLE_ROW_RE,
                _UL_RE, _OL_RE, _HR_RE, _IMG_RE):
        if pat.match(line):
            return False
    return True


def tokenize(md_text: str) -> List[Block]:
    lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: List[Block] = []
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]

        # YAML front matter
        if i == 0 and line.strip() == "---":
            i += 1
            while i < n and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # Editor-only HTML comments never belong in the exported manuscript.
        if _HTML_COMMENT_START_RE.match(line):
            while i < n:
                current = lines[i]
                i += 1
                if "-->" in current:
                    break
            continue

        # Fenced code block
        m = _FENCE_RE.match(line)
        if m:
            lang = m.group(1)
            code_lines: List[str] = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(Block(kind="code_block", language=lang,
                                code="\n".join(code_lines)))
            continue

        # Display math fence
        if _MATH_FENCE_RE.match(line):
            math_lines: List[str] = []
            i += 1
            while i < n and not _MATH_FENCE_RE.match(lines[i]):
                math_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(Block(kind="math_block", latex="\n".join(math_lines)))
            continue

        if _HTML_ANCHOR_RE.match(line.strip()):
            i += 1
            continue

        # ATX heading
        m = _HEADING_RE.match(line)
        if m:
            heading_text = m.group(2).strip()
            while _EMBEDDED_HEADING_PREFIX_RE.match(heading_text):
                heading_text = _EMBEDDED_HEADING_PREFIX_RE.sub("", heading_text, count=1).strip()
            blocks.append(Block(kind="heading", level=len(m.group(1)),
                                text=heading_text))
            i += 1
            continue

        # Horizontal rule
        if _HR_RE.match(line):
            blocks.append(Block(kind="hr"))
            i += 1
            continue

        # Standalone image
        m = _IMG_RE.match(line)
        if m:
            alt = re.sub(r"\\([\\\[\]])", r"\1", m.group(1))
            blocks.append(Block(kind="image", alt=alt, path=m.group(2)))
            i += 1
            continue
        if line.lstrip().startswith("!["):
            raise ValueError(f"Unsupported or malformed Markdown image: {line}")

        # Indented text block: preserve one source line -> one logical block line.
        m = _INDENTED_RE.match(line)
        if m and not _TABLE_ROW_RE.match(line):
            block_lines: List[str] = [m.group(1).rstrip()]
            i += 1
            while i < n:
                next_match = _INDENTED_RE.match(lines[i])
                if not next_match or not next_match.group(1).strip():
                    break
                block_lines.append(next_match.group(1).rstrip())
                i += 1
            blocks.append(Block(kind="indented_block", lines=block_lines))
            continue

        # Reference definition  [@key]: text...
        ref_m = re.match(r"^\[@([^\]]+)\]:\s*(.+)$", line)
        if ref_m:
            blocks.append(Block(kind="ref_def", text=ref_m.group(2).strip()))
            i += 1
            continue

        # Pipe table
        if _TABLE_ROW_RE.match(line):
            raw_rows: List[List[str]] = []
            while i < n and _TABLE_ROW_RE.match(lines[i]):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                raw_rows.append(cells)
                i += 1
            data = [r for r in raw_rows
                    if not all(re.match(r"^[-: ]+$", c) for c in r)]
            if data:
                blocks.append(Block(kind="table", header=data[0], rows=data[1:]))
            continue

        # Ordered list item
        m = _OL_RE.match(line)
        if m:
            item_lines = [m.group(3).strip()]
            i += 1
            while i < n and lines[i].strip():
                continuation = lines[i]
                leading = len(continuation) - len(continuation.lstrip(" \t"))
                if leading <= len(m.group(1)) or _OL_RE.match(continuation) or _UL_RE.match(continuation):
                    break
                item_lines.append(continuation.strip())
                i += 1
            blocks.append(Block(kind="list_item", ordered=True,
                                list_number=int(m.group(2)),
                                depth=len(m.group(1)) // 2, text=" ".join(item_lines)))
            continue

        # Unordered list item
        m = _UL_RE.match(line)
        if m:
            item_lines = [m.group(2).strip()]
            i += 1
            while i < n and lines[i].strip():
                continuation = lines[i]
                leading = len(continuation) - len(continuation.lstrip(" \t"))
                if leading <= len(m.group(1)) or _OL_RE.match(continuation) or _UL_RE.match(continuation):
                    break
                item_lines.append(continuation.strip())
                i += 1
            blocks.append(Block(kind="list_item", ordered=False,
                                depth=len(m.group(1)) // 2, text=" ".join(item_lines)))
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Paragraph
        para_lines = [line]
        i += 1
        while i < n and _is_continuation(lines[i]):
            para_lines.append(lines[i].rstrip())
            i += 1
        blocks.append(Block(kind="paragraph",
                            text=" ".join(l.rstrip() for l in para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Context helpers
# ---------------------------------------------------------------------------

def _plain_text(raw: str) -> str:
    t = re.sub(r"\*+|__?", "", raw)
    t = re.sub(r"\[@[^\]]+\]", "", t)
    t = re.sub(r"\^[^\^]+\^", "", t)
    t = re.sub(r"`[^`]+`", "", t)
    t = re.sub(r"\[[^\]]*\]\([^\)]*\)", "", t)
    return t.strip()


def _section_ctx(text: str) -> Optional[str]:
    return _SECTION_CONTEXT.get(text.strip().lower())


def _caption_style(raw_text: str) -> Optional[str]:
    plain = _plain_text(raw_text)
    for pat, key in _CAPTION_PATTERNS:
        if pat.match(plain):
            return key
    return None


def _should_include_in_toc(text: str) -> bool:
    normalized = text.strip().lower()
    if normalized in {
        "table of contents",
        "abstract",
        "keywords",
        "key words",
        "acknowledgments",
        "acknowledgements",
        "references",
        "reference",
    }:
        return False
    return True


def _collect_static_toc_entries(blocks: List[Block]) -> List[Tuple[int, str]]:
    entries: List[Tuple[int, str]] = []
    for block in blocks:
        if block.kind != "heading":
            continue
        text = block.text.strip()
        effective_level = 2 if block.level == 1 and _NUMBERED_SECTION_HEADING_RE.match(text) else block.level
        if effective_level not in {2, 3, 4}:
            continue
        if not text or not _should_include_in_toc(text):
            continue
        entries.append((effective_level, text))
    return entries


def _insert_static_toc(doc: Document, entries: List[Tuple[int, str]]) -> None:
    for level, text in entries:
        p = doc.add_paragraph(style=_S["toc"])
        if level == 3:
            p.paragraph_format.left_indent = Inches(0.32)
        elif level >= 4:
            p.paragraph_format.left_indent = Inches(0.58)
        apply_runs(p, parse_inline(text), spec_key="toc")


# ---------------------------------------------------------------------------
# Document body clear
# ---------------------------------------------------------------------------

def _clear_body(doc: Document) -> None:
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def _decorate_chemvellum_document(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == _S["abstract"]:
            _set_paragraph_shading(paragraph, _CHEMVELLUM["abstract_pale"])
            _set_paragraph_border(
                paragraph,
                "left",
                color=_CHEMVELLUM["teal"],
                size=14,
                space=8,
            )
            for edge in ("top", "right", "bottom"):
                _set_paragraph_border(
                    paragraph,
                    edge,
                    color=_CHEMVELLUM["abstract_pale"],
                    size=6,
                    space=6,
                )
            paragraph.paragraph_format.right_indent = Pt(13)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(
    md_path: Path,
    out_path: Path,
    template_path: Path,
    author: str = "",
    subject: str = "Scholarly review manuscript",
    keywords: str = "",
    layout_profile: str = _DEFAULT_LAYOUT_PROFILE,
) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    if re.search(r"\[@P\d{3}", md_text):
        raise SystemExit("[md2docx] ERROR: unresolved stable citation tokens remain in Markdown")
    if re.search(r"<!--\s*paragraph_id\s*:", md_text, re.I):
        raise SystemExit("[md2docx] ERROR: editor paragraph markers remain in Markdown")
    blocks  = tokenize(md_text)
    toc_entries = _collect_static_toc_entries(blocks)
    doc     = Document(str(template_path))
    _clear_body(doc)
    _configure_academic_document(doc, layout_profile)
    document_title = next(
        (
            block.text.strip()
            for block in blocks
            if block.kind == "heading"
            and block.level == 1
            and not _NUMBERED_SECTION_HEADING_RE.match(block.text.strip())
        ),
        md_path.stem,
    )
    properties = doc.core_properties
    properties.title = document_title
    properties.subject = subject.strip()
    properties.author = author.strip()
    properties.last_modified_by = author.strip()
    properties.keywords = keywords.strip()
    properties.comments = (
        "Generated from the review Markdown manuscript. "
        f"Layout profile: {layout_profile}."
    )
    regular_list_num_id: int | None = None
    regular_list_ordered: bool | None = None
    reference_counter = 0

    ctx: str           = "body"
    front_matter: bool = False
    inserted_toc_heading = False
    saw_toc_heading = False
    skipping_source_toc = False
    missing_images: List[str] = []
    expected_images = sum(1 for block in blocks if block.kind == "image")
    inserted_images = 0
    lead_body_next = True
    journal_body_started = False
    journal_columns = 1
    pending_table_caption: str | None = None
    keep_next_body_together = False

    def insert_toc_once() -> None:
        nonlocal inserted_toc_heading
        if inserted_toc_heading:
            return
        _para(doc, "body", "h2", "Table of Contents", force_bold=True)
        _insert_static_toc(doc, toc_entries)
        inserted_toc_heading = True

    for block in blocks:

        if block.kind != "list_item":
            regular_list_num_id = None
            regular_list_ordered = None
        if pending_table_caption is not None and block.kind != "table":
            _para(doc, "table_title", "table_title", pending_table_caption)
            pending_table_caption = None

        if block.kind == "heading":
            plain_heading = block.text.strip().lower()
            numbered_h1_section = block.level == 1 and _NUMBERED_SECTION_HEADING_RE.match(block.text.strip())
            if plain_heading == "table of contents":
                saw_toc_heading = True
                skipping_source_toc = True
                insert_toc_once()
                continue
            skipping_source_toc = False
            effective_level = 2 if numbered_h1_section else block.level
            style_key, spec_key = _HEADING_FORMAT.get(effective_level, ("body", "body"))
            new_ctx = _section_ctx(block.text)
            if (
                layout_profile == "chemvellum_journal"
                and not journal_body_started
                and effective_level >= 2
                and new_ctx not in {"abstract", "keywords", "acks", "supporting", "references"}
            ):
                _add_chemvellum_section(doc, 2, WD_SECTION.CONTINUOUS)
                journal_body_started = True
                journal_columns = 2
            ctx = new_ctx if new_ctx else "body"
            if block.level == 1 and not numbered_h1_section:
                front_matter = True
            elif effective_level >= 2:
                front_matter = False
            _para(doc, style_key, spec_key, block.text)
            lead_body_next = True

        elif block.kind == "paragraph":
            text  = block.text.strip()
            plain = _plain_text(text)

            if re.match(r"^\*\*keywords:?\*\*", text, re.I):
                ctx = "keywords"
                _para(doc, "keywords", "keywords", text)
                continue

            # Bold-only section label  e.g. **Abstract**
            new_ctx = _section_ctx(plain)
            if new_ctx:
                skipping_source_toc = False
                if new_ctx in {"abstract", "keywords"} and not inserted_toc_heading and not saw_toc_heading:
                    insert_toc_once()
                ctx = new_ctx
                _para(doc, "body", "body", text, force_bold=True)
                lead_body_next = True
                continue
            if skipping_source_toc:
                continue

            # Front matter: author / affiliation
            if front_matter and ctx == "body":
                if _AFFIL_START.match(text):
                    _para(doc, "address", "address", text)
                else:
                    _para(doc, "author", "author", text)
                continue

            if ctx != "body":
                spec = ctx if ctx in _FONT_SPEC else "body"
                _para(doc, ctx, spec, text)
            else:
                cap = _caption_style(text)
                if (
                    cap == "table_title"
                    and layout_profile == "chemvellum_journal"
                    and journal_body_started
                ):
                    pending_table_caption = text
                    lead_body_next = True
                    continue
                key = cap if cap else ("body_lead" if lead_body_next else "body")
                paragraph = _para(doc, key, key, text)
                if keep_next_body_together and cap is None:
                    paragraph.paragraph_format.keep_together = True
                    keep_next_body_together = False
                lead_body_next = bool(cap)

        elif block.kind == "indented_block":
            if skipping_source_toc:
                continue
            for raw_line in block.lines:
                text = raw_line.strip()
                if not text:
                    continue
                cap = _caption_style(text)
                if cap:
                    key = cap
                elif ctx == "references":
                    key = "references"
                else:
                    key = "body_lead" if lead_body_next else "body"
                spec = key if key in _FONT_SPEC else "body"
                _para(doc, key, spec, text)
                lead_body_next = bool(cap)
        elif block.kind == "ref_def":
            if skipping_source_toc:
                continue
            reference_counter += 1
            _reference_para(
                doc,
                reference_counter,
                block.text,
                layout_profile=layout_profile,
            )

        elif block.kind == "list_item":
            if skipping_source_toc:
                continue
            if ctx == "references":
                reference_counter = block.list_number or (reference_counter + 1)
                _reference_para(
                    doc,
                    reference_counter,
                    block.text,
                    layout_profile=layout_profile,
                )
            else:
                if regular_list_num_id is None or regular_list_ordered != block.ordered:
                    regular_list_num_id = _create_numbering_definition(
                        doc,
                        ordered=block.ordered,
                        start_at=block.list_number or 1,
                    )
                    regular_list_ordered = block.ordered
                p = _para(doc, "list", "list", block.text)
                _apply_numbering(p, regular_list_num_id, block.depth)
            lead_body_next = False

        elif block.kind == "code_block":
            if skipping_source_toc:
                continue
            p  = doc.add_paragraph(style=_S["body_lead"])
            wr = p.add_run(block.code)
            _set_word_run_font(wr, "Courier New", 9)
            lead_body_next = True

        elif block.kind == "math_block":
            if skipping_source_toc:
                continue
            p = doc.add_paragraph(style=_S["body_lead"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_together = True
            _apply_math(p, block.latex)
            lead_body_next = True
        elif block.kind == "table":
            if skipping_source_toc:
                continue
            resume_columns = journal_columns
            if layout_profile == "chemvellum_journal" and journal_body_started and journal_columns != 1:
                _add_chemvellum_section(doc, 1)
                journal_columns = 1
            if pending_table_caption is not None:
                _para(doc, "table_title", "table_title", pending_table_caption)
                pending_table_caption = None
            _add_table(
                doc,
                block.header,
                block.rows,
                width_inches=(
                    _CHEMVELLUM["content_width"]
                    if layout_profile == "chemvellum_journal"
                    else 6.5
                ),
                layout_profile=layout_profile,
            )
            if layout_profile == "chemvellum_journal" and journal_body_started and resume_columns == 2:
                _add_chemvellum_section(doc, 2)
                doc.paragraphs[-1].paragraph_format.space_before = Pt(4)
                journal_columns = 2
            lead_body_next = True
            keep_next_body_together = True
        elif block.kind == "image":
            if skipping_source_toc:
                continue
            img_path = Path(block.path)
            if not img_path.is_absolute():
                img_path = md_path.parent / img_path
            if img_path.exists():
                resume_columns = journal_columns
                if layout_profile == "chemvellum_journal" and journal_body_started and journal_columns != 1:
                    _add_chemvellum_section(doc, 1)
                    journal_columns = 1
                p = doc.add_paragraph(style=_S["figure_display"])
                caption_key = _caption_style(block.alt) or "figure"
                max_width, max_height = _figure_bounds_for_caption(
                    caption_key,
                    _usable_page_width_inches(doc, layout_profile),
                )
                if layout_profile == "chemvellum_journal":
                    max_height = min(max_height, 3.7)
                figure_width, figure_height = _bounded_figure_size(
                    img_path, max_width, max_height
                )
                picture_kwargs = {"width": Inches(figure_width)}
                if figure_height is not None:
                    picture_kwargs["height"] = Inches(figure_height)
                p.add_run().add_picture(str(img_path), **picture_kwargs)
                inserted_images += 1
                if block.alt:
                    _para(doc, caption_key, caption_key, block.alt)
                if layout_profile == "chemvellum_journal" and journal_body_started and resume_columns == 2:
                    _add_chemvellum_section(doc, 2)
                    journal_columns = 2
                lead_body_next = True
            else:
                missing_images.append(str(img_path))

        elif block.kind == "hr":
            # Horizontal rules in review Markdown are section separators, not
            # desired visual borders in the final DOCX.
            continue

    if missing_images:
        raise SystemExit("[md2docx] ERROR: missing images: " + ", ".join(missing_images))
    if pending_table_caption is not None:
        _para(doc, "table_title", "table_title", pending_table_caption)
    if inserted_images != expected_images:
        raise SystemExit(
            "[md2docx] ERROR: embedded "
            f"{inserted_images} of {expected_images} parsed Markdown images"
        )
    if layout_profile == "chemvellum_journal":
        if journal_body_started and journal_columns == 2:
            # A final continuous one-column section balances the last pair of
            # reference columns without adding a cover-like blank page.
            _add_chemvellum_section(doc, 1)
        _decorate_chemvellum_document(doc)
        _configure_chemvellum_header_footer(doc)
    else:
        _configure_footer(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(_writable_path(out_path))
    print(f"[md2docx] Saved -> {out_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

_DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent / "review_template.docx"
_WINDOWS_RESERVED_STEMS = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{index}" for index in range(1, 10)),
    *(f"LPT{index}" for index in range(1, 10)),
}


def manuscript_title(md_path: Path) -> str:
    """Return the first level-one Markdown heading."""
    for line in md_path.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^\s*#\s+(.+?)\s*$", line)
        if match:
            return match.group(1).strip()
    return ""


def document_filename_stem(title: str, fallback: str = "review") -> str:
    """Make a readable Windows-safe filename while preserving the paper title."""
    stem = re.sub(r"[<>:\"/\\|?*\x00-\x1f]+", " - ", title)
    stem = re.sub(r"\s+", " ", stem).strip(" .-")
    if not stem:
        stem = fallback.strip(" .") or "review"
    if stem.upper() in _WINDOWS_RESERVED_STEMS:
        stem = f"{stem} manuscript"
    return stem[:180].rstrip(" .-")


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="md2docx",
        description="Convert Markdown to DOCX using review_template.docx styles.",
    )
    p.add_argument("--input",    required=True, metavar="MD",   help="Input .md file")
    p.add_argument(
        "--output",
        metavar="DOCX",
        help="Output .docx file. Omit to use the manuscript title as the filename.",
    )
    p.add_argument(
        "--output-dir",
        metavar="DIR",
        help="Directory for the title-named DOCX when --output is omitted.",
    )
    p.add_argument("--template", default=str(_DEFAULT_TEMPLATE), metavar="DOCX",
                   help=f"Word template (default: {_DEFAULT_TEMPLATE})")
    p.add_argument("--author", default="", help="Document author metadata; empty removes stale template authors")
    p.add_argument("--subject", default="Scholarly review manuscript", help="Document subject metadata")
    p.add_argument("--keywords", default="", help="Comma- or semicolon-separated document keywords")
    p.add_argument(
        "--layout-profile",
        choices=_LAYOUT_PROFILES,
        default=_DEFAULT_LAYOUT_PROFILE,
        help=(
            "Word layout profile (default: professional_single; "
            "chemvellum_journal adds the branded two-column editorial system; "
            "legacy_report preserves the previous report-like spacing)."
        ),
    )
    return p


def main() -> None:
    args          = _build_parser().parse_args()
    md_path       = Path(args.input).resolve()
    template_path = Path(args.template).resolve()

    if not md_path.exists():
        raise SystemExit(f"[md2docx] ERROR: Input not found: {md_path}")
    if not template_path.exists():
        raise SystemExit(f"[md2docx] ERROR: Template not found: {template_path}")
    if not _LATEX_OK:
        print("[md2docx] INFO: latex2word unavailable; using built-in deterministic script formatting")
    if args.output:
        out_path = Path(args.output).resolve()
    else:
        output_dir = Path(args.output_dir).resolve() if args.output_dir else md_path.parent
        title = manuscript_title(md_path)
        out_path = output_dir / f"{document_filename_stem(title, md_path.stem)}.docx"

    convert(
        md_path,
        out_path,
        template_path,
        author=args.author,
        subject=args.subject,
        keywords=args.keywords,
        layout_profile=args.layout_profile,
    )


if __name__ == "__main__":
    main()
