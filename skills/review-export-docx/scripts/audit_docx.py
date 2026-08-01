#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


LAYOUT_PROFILE_SPECS = {
    "chemvellum_journal": {
        "styles": {
            "Review Title": {"font": "Georgia", "size": 21.5},
            "Review Body": {"font": "Cambria", "size": 9.5},
            "Review Body Lead": {"font": "Cambria", "size": 9.5},
            "Review List": {"font": "Cambria", "size": 9.0},
            "Review Heading 1": {"font": "Arial", "size": 11.5},
            "Review Heading 2": {"font": "Arial", "size": 9.5},
            "Review Heading 3": {"font": "Arial", "size": 9.0},
            "Review Abstract": {"font": "Cambria", "size": 10.0},
            "Review Reference": {"font": "Cambria", "size": 7.5},
            "Review Figure Caption": {"font": "Cambria", "size": 8.0},
            "Review Table Caption": {"font": "Cambria", "size": 8.0},
            "Review Table Header": {"font": "Arial", "size": 7.5},
            "Review Table Body": {"font": "Cambria", "size": 7.5},
        },
        "paragraphs": {
            "Review Body": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY, "line_spacing": 1.04, "after": 1.5, "first_line": 9.0},
            "Review Body Lead": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY, "line_spacing": 1.04, "after": 1.5, "first_line": 9.0},
            "Review List": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.02, "after": 1.5, "first_line": 0.0},
            "Review Reference": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.0, "after": 1.0, "first_line": -18.0, "left_indent": 18.0},
            "Review Figure Caption": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.04, "after": 6.0, "first_line": 0.0},
        },
        "geometry": (8.5, 11.0, 0.62, 0.65, 0.65, 0.65),
        "table_width_dxa": 10368,
        "table_indent_dxa": 120,
        "requires_two_columns": True,
        "wordmark": "CHEMVELLUM",
    },
    "professional_single": {
        "styles": {
            "Review Title": {"font": "Times New Roman", "size": 18.0},
            "Review Body": {"font": "Times New Roman", "size": 11.0},
            "Review Body Lead": {"font": "Times New Roman", "size": 11.0},
            "Review List": {"font": "Times New Roman", "size": 11.0},
            "Review Heading 1": {"font": "Times New Roman", "size": 13.5},
            "Review Heading 2": {"font": "Times New Roman", "size": 11.5},
            "Review Heading 3": {"font": "Times New Roman", "size": 10.5},
            "Review Abstract": {"font": "Times New Roman", "size": 10.5},
            "Review Reference": {"font": "Times New Roman", "size": 9.0},
            "Review Figure Caption": {"font": "Times New Roman", "size": 9.0},
            "Review Table Caption": {"font": "Times New Roman", "size": 9.0},
            "Review Table Body": {"font": "Times New Roman", "size": 9.0},
        },
        "paragraphs": {
            "Review Body": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY, "line_spacing": 1.15, "after": 0.0, "first_line": 11.0},
            "Review Body Lead": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY, "line_spacing": 1.15, "after": 0.0, "first_line": 0.0},
            "Review List": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.08, "after": 2.0, "first_line": 0.0},
            "Review Reference": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.0, "after": 2.0, "first_line": -36.0, "left_indent": 36.0},
            "Review Figure Caption": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.0, "after": 6.0, "first_line": 0.0},
        },
        "geometry": (8.5, 11.0, 1.0, 1.0, 1.0, 1.0),
        "table_width_dxa": 9360,
        "table_indent_dxa": 120,
    },
    "legacy_report": {
        "styles": {
            "Review Title": {"font": "Times New Roman", "size": 18.0},
            "Review Body": {"font": "Times New Roman", "size": 12.0},
            "Review Heading 1": {"font": "Times New Roman", "size": 14.0},
            "Review Heading 2": {"font": "Times New Roman", "size": 12.0},
            "Review Heading 3": {"font": "Times New Roman", "size": 11.0},
            "Review Abstract": {"font": "Times New Roman", "size": 11.0},
            "Review Reference": {"font": "Times New Roman", "size": 10.0},
            "Review Figure Caption": {"font": "Times New Roman", "size": 10.0},
        },
        "paragraphs": {
            "Review Body": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.5, "after": 6.0, "first_line": 0.0},
        },
        "geometry": (8.5, 11.0, 1.0, 1.0, 1.0, 1.0),
        "table_width_dxa": 9360,
        "table_indent_dxa": 120,
    },
}

PARENTHESIZED_COMPLEX_RE = (
    r"(?:[A-Z][a-z]?\d*)+(?:\([A-Za-z][A-Za-z0-9]*\)\d*)+(?:[A-Z][a-z]?\d*)*"
)
CHEMICAL_TOKEN_RE = re.compile(
    rf"(?<![A-Za-z0-9])(?:{PARENTHESIZED_COMPLEX_RE}|(?:[A-Z][a-z]?\d*){{2,}})(?![A-Za-z0-9])"
)


def close(value: float | None, expected: float, tolerance: float = 0.02) -> bool:
    return value is not None and abs(value - expected) <= tolerance


def style_font_name(style) -> str | None:
    if style.font.name:
        return style.font.name
    rpr = style.element.rPr
    if rpr is not None and rpr.rFonts is not None:
        return rpr.rFonts.get(qn("w:ascii"))
    return None


def _length_pt(value) -> float:
    return value.pt if value is not None else 0.0


def detect_layout_profile(document: Document) -> str:
    comments = str(document.core_properties.comments or "")
    match = re.search(r"Layout profile:\s*([a-z_]+)", comments, re.I)
    if match and match.group(1).lower() in LAYOUT_PROFILE_SPECS:
        return match.group(1).lower()
    try:
        body = document.styles["Review Body"]
    except KeyError:
        return "legacy_report"
    size = body.font.size.pt if body.font.size else None
    if close(size, 11.0) and body.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "professional_single"
    return "legacy_report"


def effective_keep_with_next(paragraph) -> bool:
    direct = paragraph.paragraph_format.keep_with_next
    if direct is not None:
        return bool(direct)
    style = paragraph.style
    while style is not None:
        inherited = style.paragraph_format.keep_with_next
        if inherited is not None:
            return bool(inherited)
        style = style.base_style
    return False


def markdown_image_count(path: Path | None) -> int | None:
    if not path:
        return None
    text = path.read_text(encoding="utf-8")
    return len(re.findall(r"!\[(?:\\.|[^\]])*\]\([^)]+\)", text))


def markdown_title(path: Path | None) -> str | None:
    if not path:
        return None
    text = path.read_text(encoding="utf-8")
    match = re.search(r"(?m)^#\s+(.+?)\s*$", text)
    return match.group(1).strip() if match else None


def unformatted_formula_tokens(document: Document) -> list[dict[str, object]]:
    issues: list[dict[str, object]] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = "".join(run.text for run in paragraph.runs)
        script_flags = [
            bool(run.font.subscript)
            for run in paragraph.runs
            for _ in run.text
        ]
        if len(script_flags) != len(text):
            continue
        for match in CHEMICAL_TOKEN_RE.finditer(text):
            token = match.group(0)
            digit_positions = [
                match.start() + offset
                for offset, char in enumerate(token)
                if char.isdigit()
            ]
            if digit_positions and any(not script_flags[position] for position in digit_positions):
                issues.append({"paragraph": paragraph_index, "token": token})
    return issues


def audit(
    docx_path: Path,
    markdown_path: Path | None,
    render_qa: str = "not_run",
    layout_profile: str = "auto",
) -> dict:
    document = Document(docx_path)
    blockers: list[str] = []
    warnings: list[str] = []
    detected_layout_profile = detect_layout_profile(document)
    selected_layout_profile = (
        detected_layout_profile if layout_profile == "auto" else layout_profile
    )
    if selected_layout_profile not in LAYOUT_PROFILE_SPECS:
        raise ValueError(f"unknown layout profile: {selected_layout_profile}")
    if layout_profile != "auto" and detected_layout_profile != selected_layout_profile:
        blockers.append(
            "document layout profile metadata/inference is "
            f"{detected_layout_profile}, expected {selected_layout_profile}"
        )
    profile_spec = LAYOUT_PROFILE_SPECS[selected_layout_profile]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    properties = document.core_properties
    expected_title = markdown_title(markdown_path)
    stale_template_titles = {"chemvellum review template"}
    stale_template_authors = {"chemvellum"}
    if not str(properties.title or "").strip():
        blockers.append("document title metadata is empty")
    elif str(properties.title).strip().lower() in stale_template_titles:
        blockers.append("document title metadata still contains the template title")
    elif expected_title and str(properties.title).strip() != expected_title:
        blockers.append("document title metadata does not match the Markdown title")
    if not str(properties.subject or "").strip():
        blockers.append("document subject metadata is empty")
    if str(properties.author or "").strip().lower() in stale_template_authors:
        blockers.append("document author metadata still contains the template author")
    if str(properties.last_modified_by or "").strip().lower() in stale_template_authors:
        blockers.append("document last-modified-by metadata still contains the template author")

    if "<!-- paragraph_id:" in text:
        blockers.append("editor paragraph markers are visible")
    if re.search(r"\[@P\d{3}", text):
        blockers.append("stable citation tokens are visible")
    if re.search(r"\[P\d{3}\]", text):
        blockers.append("internal paper IDs are visible")
    if re.search(r"\\(?:mathrm|mathbf|mathsf|ce)\b|_\s*\{|\^\s*\{", text):
        blockers.append("raw LaTeX commands are visible")

    for name, expected in profile_spec["styles"].items():
        try:
            style = document.styles[name]
        except KeyError:
            blockers.append(f"required style missing: {name}")
            continue
        if style_font_name(style) != expected["font"]:
            blockers.append(f"{name} font is not {expected['font']}")
        size = style.font.size.pt if style.font.size else None
        if not close(size, expected["size"]):
            blockers.append(f"{name} size is {size}, expected {expected['size']}")

    for name, expected in profile_spec["paragraphs"].items():
        try:
            paragraph = document.styles[name].paragraph_format
        except KeyError:
            continue
        if paragraph.alignment != expected["alignment"]:
            blockers.append(
                f"{name} alignment is {paragraph.alignment}, expected {expected['alignment']}"
            )
        spacing = float(paragraph.line_spacing) if paragraph.line_spacing else None
        if not close(spacing, expected["line_spacing"]):
            blockers.append(
                f"{name} line spacing is {spacing}, expected {expected['line_spacing']}"
            )
        after = _length_pt(paragraph.space_after)
        if not close(after, expected["after"]):
            blockers.append(
                f"{name} paragraph spacing after is {after}, expected {expected['after']} pt"
            )
        first_line = _length_pt(paragraph.first_line_indent)
        if not close(first_line, expected["first_line"]):
            blockers.append(
                f"{name} first-line indent is {first_line}, expected {expected['first_line']} pt"
            )
        if "left_indent" in expected:
            left_indent = _length_pt(paragraph.left_indent)
            if not close(left_indent, expected["left_indent"]):
                blockers.append(
                    f"{name} left indent is {left_indent}, expected {expected['left_indent']} pt"
                )

    if len(document.sections) != 1 and selected_layout_profile != "chemvellum_journal":
        warnings.append(f"document has {len(document.sections)} sections")
    expected_geometry = profile_spec["geometry"]
    two_column_sections = 0
    for index, section in enumerate(document.sections, start=1):
        geometry = (
            section.page_width.inches,
            section.page_height.inches,
            section.top_margin.inches,
            section.bottom_margin.inches,
            section.left_margin.inches,
            section.right_margin.inches,
        )
        if any(not close(actual, wanted) for actual, wanted in zip(geometry, expected_geometry)):
            blockers.append(
                f"section {index} page geometry is {geometry}, expected {expected_geometry}"
            )
        cols = section._sectPr.find(qn("w:cols"))
        column_count = int(cols.get(qn("w:num")) or 1) if cols is not None else 1
        two_column_sections += int(column_count == 2)
    if profile_spec.get("requires_two_columns") and two_column_sections == 0:
        blockers.append("ChemVellum journal layout has no two-column body section")
    expected_wordmark = profile_spec.get("wordmark")
    if expected_wordmark:
        header_text = "".join(
            paragraph.text
            for section in document.sections
            for paragraph in section.header.paragraphs
        )
        if expected_wordmark not in header_text:
            blockers.append(f"running header is missing the {expected_wordmark} wordmark")

    drawing_count = len(document.inline_shapes)
    expected_images = markdown_image_count(markdown_path)
    if expected_images is not None and drawing_count != expected_images:
        blockers.append(f"DOCX contains {drawing_count} images; Markdown contains {expected_images}")
    abstract_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip().lower() == "abstract"
        ),
        None,
    )
    introduction_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if re.fullmatch(
                r"(?:\d+[.)]?\s*)?introduction",
                paragraph.text.strip(),
                re.I,
            )
        ),
        None,
    )
    if abstract_index is not None and introduction_index is not None and any(
        paragraph._p.xpath(".//w:drawing")
        for paragraph in document.paragraphs[abstract_index + 1 : introduction_index]
    ):
        blockers.append("a figure appears inside the Abstract block")

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        if not paragraph._p.xpath(".//w:drawing"):
            continue
        keep_next = effective_keep_with_next(paragraph)
        if not keep_next:
            blockers.append(f"figure paragraph {paragraph_index + 1} is not kept with its caption")
        next_paragraph = document.paragraphs[paragraph_index + 1] if paragraph_index + 1 < len(document.paragraphs) else None
        if next_paragraph is None or next_paragraph.style.name != "Review Figure Caption":
            blockers.append(f"figure paragraph {paragraph_index + 1} is not followed by a figure caption")

    for table_index, table in enumerate(document.tables, start=1):
        expected_table_width = str(profile_spec["table_width_dxa"])
        expected_table_indent = str(profile_spec["table_indent_dxa"])
        widths = table._tbl.tblPr.findall(qn("w:tblW"))
        if len(widths) != 1 or widths[0].get(qn("w:w")) != expected_table_width or widths[0].get(qn("w:type")) != "dxa":
            blockers.append(
                f"table {table_index} does not have exact {expected_table_width} DXA width"
            )
        indents = table._tbl.tblPr.findall(qn("w:tblInd"))
        if len(indents) != 1 or indents[0].get(qn("w:w")) != expected_table_indent:
            blockers.append(
                f"table {table_index} does not have {expected_table_indent} DXA indent"
            )
        grid = [int(column.get(qn("w:w")) or 0) for column in table._tbl.tblGrid.findall(qn("w:gridCol"))]
        if sum(grid) != int(expected_table_width):
            blockers.append(
                f"table {table_index} grid width sums to {sum(grid)}, expected {expected_table_width}"
            )
        for row_index, row in enumerate(table.rows, start=1):
            cell_widths = [int(cell._tc.tcPr.tcW.get(qn("w:w")) or 0) for cell in row.cells]
            if cell_widths != grid:
                blockers.append(
                    f"table {table_index} row {row_index} cell widths {cell_widths} do not match grid {grid}"
                )
            row_properties = row._tr.get_or_add_trPr()
            if row_properties.find(qn("w:cantSplit")) is None:
                blockers.append(f"table {table_index} row {row_index} can split across pages")
        if table.rows:
            header_properties = table.rows[0]._tr.get_or_add_trPr()
            if header_properties.find(qn("w:tblHeader")) is None:
                blockers.append(f"table {table_index} header row is not marked to repeat")
        if len(table.rows) >= 3:
            protected_rows = {0, 1, len(table.rows) - 2}
            for row_index in sorted(protected_rows):
                if not all(
                    paragraph.paragraph_format.keep_with_next
                    for cell in table.rows[row_index].cells
                    for paragraph in cell.paragraphs
                ):
                    blockers.append(
                        f"table {table_index} row {row_index + 1} lacks orphan protection"
                    )

    subscript_runs = 0
    superscript_runs = 0
    numbered_paragraphs = 0
    visible_reference_labels = 0
    direct_font_or_size_runs = 0
    for paragraph in document.paragraphs:
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            numbered_paragraphs += 1
        if paragraph.style.name == "Review Reference" and re.match(
            r"^\s*\d+\.\s+\S", paragraph.text
        ):
            visible_reference_labels += 1
        for run in paragraph.runs:
            subscript_runs += int(bool(run.font.subscript))
            superscript_runs += int(bool(run.font.superscript))
            direct_font_or_size_runs += int(
                bool(run.text.strip())
                and (run.font.name is not None or run.font.size is not None)
            )

    formula_issues = unformatted_formula_tokens(document)
    if formula_issues:
        blockers.append(
            "chemical formula digits remain at baseline: "
            + ", ".join(
                f"{item['token']} (paragraph {item['paragraph']})"
                for item in formula_issues[:10]
            )
        )
    if render_qa == "not_run":
        warnings.append("visual render QA has not been performed")
    elif render_qa == "unavailable":
        warnings.append("visual render QA was unavailable")

    reference_heading_index = next(
        (index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip().lower() == "references"),
        None,
    )
    if reference_heading_index is None:
        blockers.append("References heading is missing")
    elif not any(
        (
            paragraph._p.pPr is not None
            and paragraph._p.pPr.numPr is not None
        )
        or re.match(r"^\d+\.\s+\S", paragraph.text.strip())
        for paragraph in document.paragraphs[reference_heading_index + 1 :]
    ):
        blockers.append("References are not visibly numbered")

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        ppr = paragraph._p.pPr
        if ppr is None or ppr.numPr is None:
            continue
        expected_style = (
            "Review Reference"
            if reference_heading_index is not None and paragraph_index > reference_heading_index
            else "Review List"
        )
        if selected_layout_profile in {"professional_single", "chemvellum_journal"} and paragraph.style.name != expected_style:
            blockers.append(
                f"numbered paragraph {paragraph_index + 1} uses {paragraph.style.name}, "
                f"expected {expected_style}"
            )

    return {
        "docx_path": str(docx_path),
        "markdown_path": str(markdown_path) if markdown_path else None,
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_image_count": drawing_count,
        "markdown_image_count": expected_images,
        "metadata": {
            "title": properties.title,
            "subject": properties.subject,
            "author": properties.author,
            "last_modified_by": properties.last_modified_by,
            "keywords": properties.keywords,
        },
        "numbered_paragraph_count": numbered_paragraphs,
        "visible_reference_label_count": visible_reference_labels,
        "subscript_run_count": subscript_runs,
        "superscript_run_count": superscript_runs,
        "unformatted_formula_tokens": formula_issues,
        "render_qa": render_qa,
        "layout_profile": selected_layout_profile,
        "detected_layout_profile": detected_layout_profile,
        "style_observations": {
            "direct_font_or_size_run_count": direct_font_or_size_runs,
            "note": "informational only; semantic emphasis and scripts may remain direct formatting",
        },
        "blocking_issues": blockers,
        "warnings": warnings,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Audit review DOCX structure and academic style tokens.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--markdown", type=Path)
    parser.add_argument("--output-json", type=Path)
    parser.add_argument("--render-qa", choices=("passed", "unavailable", "not_run"), default="not_run")
    parser.add_argument(
        "--layout-profile",
        choices=("auto", *LAYOUT_PROFILE_SPECS),
        default="auto",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    report = audit(
        args.input.resolve(),
        args.markdown.resolve() if args.markdown else None,
        args.render_qa,
        args.layout_profile,
    )
    rendered = json.dumps(report, ensure_ascii=False, indent=2)
    if args.output_json:
        args.output_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_json.write_text(rendered + "\n", encoding="utf-8")
    print(rendered)
    return 1 if report["blocking_issues"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
