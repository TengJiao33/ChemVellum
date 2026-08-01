from __future__ import annotations

import importlib.util
import base64
import json
import os
import sys
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace


REPO = Path(__file__).resolve().parents[1]
INSPECT = REPO / "skills" / "review-writing-tools" / "scripts" / "inspect_review.py"
SEARCH = REPO / "skills" / "review-writing-tools" / "scripts" / "search_library.py"
MERGE = REPO / "skills" / "review-citation-assets" / "scripts" / "merge_citations.py"
INSERT = REPO / "skills" / "review-citation-assets" / "scripts" / "insert_assets.py"
RENDER = REPO / "skills" / "review-export-docx" / "scripts" / "render_docx.py"
MD2DOCX = REPO / "skills" / "review-export-docx" / "scripts" / "md2docx.py"
AUDIT_DOCX = REPO / "skills" / "review-export-docx" / "scripts" / "audit_docx.py"


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def write_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


class ReviewToolTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.metadata_dir = (
            self.root / "review-library" / "metadata" / "papers"
        )
        self.metadata_dir.mkdir(parents=True)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def add_paper(
        self,
        paper_id: str,
        *,
        title: str = "A useful paper",
        authors: object | None = None,
        journal: str = "Journal of Useful Results",
        year: int | None = 2024,
        doi: str | None = "10.1000/example",
        full_text: str = "The method improves selectivity under bounded conditions.",
    ) -> Path:
        source = self.root / "sources" / f"{paper_id}.md"
        source.parent.mkdir(parents=True, exist_ok=True)
        source.write_text(full_text, encoding="utf-8")
        payload = {
            "paper_id": paper_id,
            "title": {"value": title},
            "authors": authors
            if authors is not None
            else [{"given": "Ada", "family": "Lovelace"}],
            "journal": {"value": journal} if journal else {"value": ""},
            "year": {"value": year},
            "doi": {"value": doi},
            "abstract": {"value": full_text},
            "source_paths": {"markdown": str(source)},
        }
        write_json(self.metadata_dir / f"{paper_id}.metadata.json", payload)
        return source

    def test_word_range_is_a_source_count_advisory(self) -> None:
        module = load_module("inspect_review_range", INSPECT)
        few = module.recommended_range("comprehensive", 5)
        many = module.recommended_range("comprehensive", 35)
        self.assertEqual(few["center"], 4000)
        self.assertGreater(many["center"], few["center"])
        self.assertEqual(many["basis"], "cited_papers_with_local_full_text")

    def test_inspector_does_not_gate_a_short_manuscript(self) -> None:
        module = load_module("inspect_review_short", INSPECT)
        self.add_paper("P001")
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "# Review\n\n## Discussion\n\nA bounded result is reported [@P001].\n",
            encoding="utf-8",
        )
        report = module.inspect(self.root, manuscript, "comprehensive")
        self.assertEqual(report["mechanical_errors"], [])
        self.assertNotIn("word_advisory", report)
        self.assertEqual(report["report_type"], "advisory_snapshot")
        rendered = json.dumps(report).lower()
        self.assertNotIn('"passed"', rendered)
        self.assertNotIn('"release_ready"', rendered)

        with_advisory = module.inspect(
            self.root,
            manuscript,
            "comprehensive",
            include_word_advisory=True,
        )
        self.assertEqual(with_advisory["word_advisory"]["low"], 3200)
        self.assertNotIn("length_observation", with_advisory)

    def test_inspector_reports_broken_local_references_only(self) -> None:
        module = load_module("inspect_review_broken", INSPECT)
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "# Review\n\nUnknown evidence [@P999].\n\n![Figure 1](missing.png)\n",
            encoding="utf-8",
        )
        report = module.inspect(self.root, manuscript, "focused")
        self.assertIn("unknown stable citation: P999", report["mechanical_errors"])
        self.assertIn("missing local image: missing.png", report["mechanical_errors"])

    def test_library_search_uses_full_text(self) -> None:
        module = load_module("search_library_fulltext", SEARCH)
        self.add_paper(
            "P010",
            title="General method",
            full_text="A crossover experiment constrains the proposed radical mechanism.",
        )
        rows = module.search(
            self.root,
            "crossover radical mechanism",
            limit=10,
            metadata_only=False,
        )
        self.assertEqual(rows[0]["paper_id"], "P010")
        self.assertIn("full_text", rows[0]["matched_fields"])

    def test_relative_library_paths_survive_checkout_relocation(self) -> None:
        search_module = load_module("search_library_relative", SEARCH)
        inspect_module = load_module("inspect_review_relative", INSPECT)
        source = self.add_paper(
            "P011",
            title="Portable local source",
            full_text="A relocation control preserves access to the full text.",
        )
        metadata_path = self.metadata_dir / "P011.metadata.json"
        metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
        metadata["source_paths"]["markdown"] = source.relative_to(self.root).as_posix()
        write_json(metadata_path, metadata)

        rows = search_module.search(
            self.root,
            "relocation control",
            limit=10,
            metadata_only=False,
        )
        self.assertEqual(rows[0]["paper_id"], "P011")
        self.assertIn("full_text", rows[0]["matched_fields"])

        manuscript = self.root / "manuscript.md"
        manuscript.write_text("# Review\n\nPortable evidence [@P011].\n", encoding="utf-8")
        report = inspect_module.inspect(self.root, manuscript, "focused")
        self.assertEqual(report["citations"][0]["local_full_text"], True)

    def test_relative_paths_do_not_escape_to_process_working_directory(self) -> None:
        search_module = load_module("search_library_root_isolation", SEARCH)
        inspect_module = load_module("inspect_review_root_isolation", INSPECT)
        insert_module = load_module("insert_assets_root_isolation", INSERT)
        source = self.add_paper(
            "P012",
            title="Root-scoped source",
            full_text="The intended checkout contains the xenotime control experiment.",
        )
        metadata_path = self.metadata_dir / "P012.metadata.json"
        metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
        relative_source = source.relative_to(self.root)
        metadata["source_paths"]["markdown"] = relative_source.as_posix()
        write_json(metadata_path, metadata)

        manifest = self.root / "asset-selection" / "asset_manifest.json"
        intended_asset = manifest.parent / "figure.png"
        intended_asset.parent.mkdir(parents=True)
        intended_asset.write_bytes(b"intended")

        original_cwd = Path.cwd()
        with tempfile.TemporaryDirectory() as shadow_name:
            shadow_root = Path(shadow_name)
            shadow_source = shadow_root / relative_source
            shadow_source.parent.mkdir(parents=True)
            shadow_source.write_text("unrelated working-directory file", encoding="utf-8")
            (shadow_root / "figure.png").write_bytes(b"shadow")
            try:
                os.chdir(shadow_root)
                rows = search_module.search(
                    self.root,
                    "xenotime control",
                    limit=10,
                    metadata_only=False,
                )
                self.assertEqual(rows[0]["source_path"], str(source.resolve()))
                self.assertEqual(
                    inspect_module.resolve_local_path(self.root, relative_source),
                    source.resolve(),
                )
                self.assertEqual(
                    insert_module.resolve_asset_path(manifest, "figure.png"),
                    intended_asset.resolve(),
                )
            finally:
                os.chdir(original_cwd)

    def test_citation_merge_numbers_first_appearance(self) -> None:
        module = load_module("merge_citations_order", MERGE)
        self.add_paper("P001", title="First paper")
        self.add_paper("P002", title="Second paper")
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "# Review\n\nSecond appears first [@P002]. "
            "Then both are compared [@P001; @P002].\n",
            encoding="utf-8",
        )
        output = self.root / "deliverables" / "review.md"
        citations = self.root / "deliverables" / "citations.json"
        report = module.merge(
            self.root,
            manuscript,
            output,
            citations,
            strict_metadata=False,
        )
        rendered = output.read_text(encoding="utf-8")
        self.assertIn("Second appears first [1]", rendered)
        self.assertIn("Then both are compared [2, 1]", rendered)
        self.assertLess(rendered.index("Second paper"), rendered.index("First paper"))
        self.assertEqual(report["unknown_paper_ids"], [])
        self.assertNotIn("{'given'", rendered)

    def test_citation_merge_reuses_a_trailing_references_heading(self) -> None:
        module = load_module("merge_citations_heading", MERGE)
        self.add_paper("P001")
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "# Review\n\nA result [@P001].\n\n## References\n",
            encoding="utf-8",
        )
        output = self.root / "review.md"
        module.merge(
            self.root,
            manuscript,
            output,
            self.root / "citations.json",
            strict_metadata=False,
        )
        rendered = output.read_text(encoding="utf-8")
        self.assertEqual(rendered.count("## References"), 1)

    def test_citation_merge_handles_escaped_image_citations_and_html_metadata(
        self,
    ) -> None:
        module = load_module("merge_citations_image_html", MERGE)
        self.add_paper(
            "P001",
            title="Activation of CO<sub>2</sub> &amp; H<sub>2</sub>",
            journal="<i>Catalysis Journal</i>",
        )
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "![Figure 1. Reproduced under CC BY 4.0 [@P001\\].](figure.png)\n",
            encoding="utf-8",
        )
        output = self.root / "review.md"
        module.merge(
            self.root,
            manuscript,
            output,
            self.root / "citations.json",
            strict_metadata=False,
        )
        rendered = output.read_text(encoding="utf-8")
        self.assertIn("CC BY 4.0 [1\\].", rendered)
        self.assertIn("Activation of CO2 & H2.", rendered)
        self.assertIn("*Catalysis Journal*.", rendered)
        self.assertNotIn("<sub>", rendered)
        self.assertNotIn("<i>", rendered)

    def test_assets_survive_citation_merge_and_docx_export(self) -> None:
        insert_module = load_module("insert_assets_e2e", INSERT)
        merge_module = load_module("merge_citations_e2e", MERGE)
        export_module = load_module("md2docx_e2e", MD2DOCX)
        audit_module = load_module("audit_docx_e2e", AUDIT_DOCX)
        self.add_paper("P001")

        source = self.root / "selected" / "figure.png"
        source.parent.mkdir(parents=True)
        source.write_bytes(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwC"
                "AAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
            )
        )
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "# Review\n\nEvidence is summarized [@P001].\n\n"
            "<!-- insert:figure-1 -->\n",
            encoding="utf-8",
        )
        manifest = self.root / "assets" / "asset_manifest.json"
        write_json(
            manifest,
            [
                {
                    "asset_id": "figure-1",
                    "kind": "figure",
                    "label": "Figure 1",
                    "path": str(source),
                    "caption": "A compact evidence map.",
                    "origin": "source_paper",
                    "source_paper_id": "P001",
                    "source_locator": "Figure 1",
                    "reuse_basis": "CC BY 4.0",
                    "attribution": "Reproduced under CC BY 4.0 [@P001].",
                }
            ],
        )
        with_assets = self.root / "manuscript.with-assets.md"
        insertion = insert_module.insert_assets(
            manuscript,
            manifest,
            with_assets,
            self.root / "assets",
        )
        self.assertEqual(insertion["mechanical_errors"], [])

        deliverables = self.root / "deliverables"
        review = deliverables / "review.md"
        merge_module.merge(
            self.root,
            with_assets,
            review,
            deliverables / "citations.json",
            strict_metadata=False,
        )
        rendered = review.read_text(encoding="utf-8")
        self.assertIn("](../assets/figure-1.png)", rendered)
        self.assertEqual(audit_module.markdown_image_count(review), 1)
        image_blocks = [
            block
            for block in export_module.tokenize(rendered)
            if block.kind == "image"
        ]
        self.assertEqual(len(image_blocks), 1)
        self.assertIn("[1]", image_blocks[0].alt)
        self.assertNotIn(r"\]", image_blocks[0].alt)

        output = deliverables / "review.docx"
        template = (
            REPO / "skills" / "review-export-docx" / "review_template.docx"
        )
        export_module.convert(review, output, template)
        from docx import Document

        self.assertEqual(len(Document(str(output)).inline_shapes), 1)

    def test_placeholder_authors_are_exposed_not_disguised(self) -> None:
        module = load_module("merge_citations_placeholder", MERGE)
        self.add_paper(
            "P003",
            authors=[{"given": "A.", "family": "Author"}],
        )
        manuscript = self.root / "manuscript.md"
        manuscript.write_text("A claim [@P003].\n", encoding="utf-8")
        output = self.root / "review.md"
        citations = self.root / "citations.json"
        report = module.merge(
            self.root,
            manuscript,
            output,
            citations,
            strict_metadata=False,
        )
        self.assertIn("[authors unavailable]", output.read_text(encoding="utf-8"))
        self.assertIn(
            "authors_missing_or_placeholder",
            report["metadata_observations"][0]["issues"],
        )

    def test_suspicious_bibliographic_metadata_is_reported(self) -> None:
        module = load_module("merge_citations_suspicious", MERGE)
        self.add_paper(
            "P005",
            title=r"Catalysis with \mathrm{CO_2}",
            authors=[
                "Axially Chiral Allenes",
                "Division of Chemistry and Chemical Engineering",
            ],
        )
        manuscript = self.root / "manuscript.md"
        manuscript.write_text("A claim [@P005].\n", encoding="utf-8")
        report = module.merge(
            self.root,
            manuscript,
            self.root / "review.md",
            self.root / "citations.json",
            strict_metadata=False,
        )
        issues = report["metadata_observations"][0]["issues"]
        self.assertIn("authors_look_title_like", issues)
        self.assertIn("authors_include_affiliation_text", issues)
        self.assertIn("title_contains_raw_latex", issues)

    def test_inspector_reports_missing_abstract_and_legacy_chemistry_markup(self) -> None:
        module = load_module("inspect_review_editorial", INSPECT)
        self.add_paper("P006")
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "# Review\n\n## Discussion\n\nCO~2~ is discussed [@P006].\n",
            encoding="utf-8",
        )
        report = module.inspect(self.root, manuscript, "comprehensive")
        observations = report["editorial_observations"]
        self.assertIn("no Abstract heading was found", observations)
        self.assertTrue(any("legacy ~...~" in item for item in observations))
        self.assertEqual(report["mechanical_errors"], [])

    def test_strict_metadata_is_optional_and_transparent(self) -> None:
        module = load_module("merge_citations_strict", MERGE)
        self.add_paper("P004", journal="", year=None, doi=None)
        manuscript = self.root / "manuscript.md"
        manuscript.write_text("A claim [@P004].\n", encoding="utf-8")
        output = self.root / "review.md"
        citations = self.root / "citations.json"
        report = module.merge(
            self.root,
            manuscript,
            output,
            citations,
            strict_metadata=True,
        )
        self.assertFalse(output.exists())
        self.assertTrue(citations.exists())
        self.assertTrue(report["metadata_observations"])

    def test_reaction_scheme_uses_the_normal_asset_inserter(self) -> None:
        module = load_module("insert_assets_scheme", INSERT)
        source = self.root / "selected" / "scheme.png"
        source.parent.mkdir(parents=True)
        source.write_bytes(b"\x89PNG\r\n\x1a\nfixture")
        manuscript = self.root / "manuscript.md"
        manuscript.write_text(
            "# Review\n\nThe reaction network is summarized below.\n\n"
            "<!-- insert:scheme-1 -->\n",
            encoding="utf-8",
        )
        manifest = self.root / "assets" / "asset_manifest.json"
        write_json(
            manifest,
            {
                "assets": [
                    {
                        "asset_id": "scheme-1",
                        "kind": "scheme",
                        "label": "Scheme 1",
                        "path": str(source),
                        "caption": "Representative reaction pathways.",
                        "origin": "source_paper",
                        "source_paper_id": "P001",
                        "source_locator": "Scheme 2, page 5",
                        "reuse_basis": "CC BY 4.0; credit line checked",
                        "attribution": "Reused under CC BY 4.0 [P001].",
                    }
                ]
            },
        )
        output = self.root / "deliverables" / "review.md"
        report = module.insert_assets(
            manuscript,
            manifest,
            output,
            output.parent / "assets",
        )
        rendered = output.read_text(encoding="utf-8")
        self.assertIn("![Scheme 1.", rendered)
        self.assertIn("[@P001\\]", rendered)
        self.assertNotIn("<!-- insert:scheme-1 -->", rendered)
        self.assertEqual(report["inserted"][0]["kind"], "scheme")
        self.assertTrue((output.parent / "assets" / "scheme-1.png").is_file())

    def test_asset_caption_does_not_repeat_its_label(self) -> None:
        module = load_module("insert_assets_label", INSERT)
        source = self.root / "selected" / "figure.png"
        source.parent.mkdir(parents=True)
        source.write_bytes(b"image")
        manuscript = self.root / "manuscript.md"
        manuscript.write_text("<!-- insert:f2 -->\n", encoding="utf-8")
        manifest = self.root / "asset_manifest.json"
        write_json(
            manifest,
            {
                "assets": [
                    {
                        "asset_id": "f2",
                        "kind": "figure",
                        "label": "Fig. 2",
                        "path": str(source),
                        "caption": "Fig. 2. Reaction scope.",
                        "origin": "original",
                    }
                ]
            },
        )
        output = self.root / "review.md"
        report = module.insert_assets(
            manuscript,
            manifest,
            output,
            self.root / "copied",
        )
        rendered = output.read_text(encoding="utf-8")
        self.assertEqual(report["mechanical_errors"], [])
        self.assertIn("![Fig. 2. Reaction scope.]", rendered)
        self.assertNotIn("Fig. 2. Fig. 2.", rendered)

    def test_source_asset_without_provenance_is_not_inserted(self) -> None:
        module = load_module("insert_assets_provenance", INSERT)
        source = self.root / "figure.png"
        source.write_bytes(b"image")
        manuscript = self.root / "manuscript.md"
        manuscript.write_text("<!-- insert:f1 -->\n", encoding="utf-8")
        manifest = self.root / "asset_manifest.json"
        write_json(
            manifest,
            {
                "assets": [
                    {
                        "asset_id": "f1",
                        "path": str(source),
                        "caption": "A figure.",
                        "origin": "source_paper",
                    }
                ]
            },
        )
        output = self.root / "out.md"
        report = module.insert_assets(
            manuscript,
            manifest,
            output,
            self.root / "copied",
        )
        self.assertFalse(output.exists())
        self.assertGreaterEqual(len(report["mechanical_errors"]), 4)

    def test_active_skill_surface_is_small_and_tool_oriented(self) -> None:
        skill_names: set[str] = set()
        for skill_file in (REPO / "skills").glob("*/SKILL.md"):
            text = skill_file.read_text(encoding="utf-8")
            match = __import__("re").search(r"^name:\s*(\S+)", text, __import__("re").M)
            self.assertIsNotNone(match)
            skill_names.add(match.group(1))
        self.assertEqual(
            skill_names,
            {
                "mineru-precise-parse-chemvellum",
                "review-citation-assets",
                "review-export-docx",
                "review-metadata-prep",
                "review-source-figure-tools",
                "review-topic-paper-discovery",
                "review-writing-tools",
            },
        )

    def test_legacy_control_plane_is_removed(self) -> None:
        legacy_paths = [
            REPO
            / "skills"
            / "review-writing-orchestrator"
            / "scripts"
            / "project_status.py",
            REPO
            / "skills"
            / "review-writing-orchestrator"
            / "scripts"
            / "review_quality_gate.py",
            REPO
            / "skills"
            / "review-writing-orchestrator"
            / "scripts"
            / "finalize_run.py",
            REPO
            / "skills"
            / "review-final-audit-release"
            / "scripts"
            / "final_audit_scan.py",
        ]
        self.assertTrue(all(not path.exists() for path in legacy_paths))

    def test_docx_scheme_is_compact_and_figures_keep_full_width(self) -> None:
        module = load_module("md2docx_bounds", MD2DOCX)
        self.assertEqual(module._figure_bounds_for_caption("scheme", 6.2), (5.25, 3.6))
        self.assertEqual(module._figure_bounds_for_caption("figure", 6.2), (6.2, 5.9))

    def test_docx_audit_rejects_unscrubbed_public_template_metadata(self) -> None:
        from docx import Document

        module = load_module("audit_docx_template_metadata", AUDIT_DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        document = Document(str(template))
        output = self.root / "unscrubbed-template.docx"
        document.save(output)

        report = module.audit(output, None)
        self.assertIn(
            "document title metadata still contains the template title",
            report["blocking_issues"],
        )
        self.assertIn(
            "document author metadata still contains the template author",
            report["blocking_issues"],
        )

    def test_docx_pairs_adjacent_subscript_and_superscript(self) -> None:
        from docx import Document

        module = load_module("md2docx_stacked_scripts", MD2DOCX)
        runs = module.parse_inline(
            "1064 mmol g_cat._^-1^ h^-1^; NO_3_^−^ reduction"
        )
        self.assertEqual(
            [run.math for run in runs if run.math],
            ["g_{cat.}^{−1}", "NO_{3}^{−}"],
        )

        document = Document()
        paragraph = document.add_paragraph()
        module.apply_runs(paragraph, runs)
        xml = paragraph._p.xml
        self.assertEqual(xml.count("<m:sSubSup>"), 2)
        self.assertIn("<m:t>cat.</m:t>", xml)
        self.assertIn("<m:t>−1</m:t>", xml)

    def test_docx_parenthesized_formula_accepts_explicit_subscript(self) -> None:
        module = load_module("md2docx_parenthesized_subscript", MD2DOCX)
        runs = module.parse_inline("Co(OH)_2_ and Ni(OH)_2_")
        self.assertEqual(
            [(run.text, run.subscript) for run in runs],
            [("Co(OH)", False), ("2", True), (" and Ni(OH)", False), ("2", True)],
        )

    def test_docx_normalizes_isotope_prefix_before_formula_scripts(self) -> None:
        module = load_module("md2docx_isotope_prefix", MD2DOCX)
        runs = module.parse_inline("^15_NO_3_^−^")
        self.assertEqual(
            [(run.text, run.superscript, run.math) for run in runs],
            [("15", True, ""), ("", False, "NO_{3}^{−}")],
        )

    def test_docx_filename_uses_windows_safe_manuscript_title(self) -> None:
        module = load_module("md2docx_title_filename", MD2DOCX)
        markdown = self.root / "manuscript.md"
        markdown.write_text(
            "# Defect Engineering in MOFs: Synthesis / Catalysis?\n",
            encoding="utf-8",
        )
        self.assertEqual(
            module.manuscript_title(markdown),
            "Defect Engineering in MOFs: Synthesis / Catalysis?",
        )
        self.assertEqual(
            module.document_filename_stem(module.manuscript_title(markdown)),
            "Defect Engineering in MOFs - Synthesis - Catalysis",
        )
        self.assertEqual(module.document_filename_stem("CON"), "CON manuscript")
        writable = module._writable_path(self.root / ("long-" + "x" * 220 + ".docx"))
        if sys.platform == "win32":
            self.assertTrue(writable.startswith("\\\\?\\"))

    def test_docx_table_uses_white_three_line_borders(self) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        module = load_module("md2docx_table", MD2DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        document = Document(str(template))
        module._clear_body(document)
        module._configure_academic_document(document)
        module._add_table_single(
            document,
            ["Method", "Yield"],
            [["Route A", "80%"], ["Route B", "75%"]],
        )
        table = document.tables[0]

        def border_value(cell, edge: str) -> str | None:
            borders = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
            element = borders.find(qn(f"w:{edge}"))
            return element.get(qn("w:val")) if element is not None else None

        self.assertEqual(border_value(table.cell(0, 0), "top"), "single")
        self.assertEqual(border_value(table.cell(0, 0), "bottom"), "single")
        self.assertEqual(border_value(table.cell(0, 0), "left"), "nil")
        self.assertEqual(border_value(table.cell(1, 0), "top"), "nil")
        self.assertEqual(border_value(table.cell(1, 0), "bottom"), "nil")
        self.assertEqual(border_value(table.cell(2, 0), "bottom"), "single")
        self.assertIsNone(table.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd")))
        header_props = table.rows[0]._tr.get_or_add_trPr()
        self.assertIsNotNone(header_props.find(qn("w:tblHeader")))
        self.assertTrue(
            all(
                row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
                for row in table.rows
            )
        )

    def test_docx_splits_six_column_tables_for_readability(self) -> None:
        from docx import Document

        module = load_module("md2docx_wide_table", MD2DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        document = Document(str(template))
        module._clear_body(document)
        module._configure_academic_document(document)
        module._add_table(
            document,
            ["Method", "Substrate", "Reagent", "Condition", "Outcome", "Scope"],
            [["Route A", "Alcohol", "Cu", "25 C", "Allene", "Broad"]],
        )
        self.assertEqual(len(document.tables), 2)
        self.assertEqual(len(document.tables[0].columns), 5)
        self.assertEqual(len(document.tables[1].columns), 2)

    def test_docx_professional_single_uses_semantic_academic_styles(self) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        module = load_module("md2docx_professional_styles", MD2DOCX)
        audit_module = load_module("audit_docx_professional_styles", AUDIT_DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        markdown = self.root / "professional.md"
        markdown.write_text(
            "# Professional Review\n\n"
            "## Discussion\n\n"
            "Lead paragraph.\n\nSecond paragraph.\n\n"
            "1. First point.\n2. Second point.\n\n"
            "After the list.\n\n"
            "## References\n\n1. First paper.\n2. Second paper.\n",
            encoding="utf-8",
        )
        output = self.root / "professional.docx"
        module.convert(markdown, output, template)
        document = Document(str(output))

        body = document.styles["Review Body"].paragraph_format
        self.assertEqual(document.styles["Review Body"].font.size.pt, 11.0)
        self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertAlmostEqual(float(body.line_spacing), 1.15, places=2)
        self.assertAlmostEqual(body.first_line_indent.pt, 11.0, places=2)
        self.assertEqual(document.styles["Review Figure Caption"].paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        by_text = {paragraph.text: paragraph for paragraph in document.paragraphs}
        self.assertEqual(by_text["Lead paragraph."].style.name, "Review Body Lead")
        self.assertEqual(by_text["Second paragraph."].style.name, "Review Body")
        self.assertEqual(by_text["First point."].style.name, "Review List")
        self.assertIsNotNone(by_text["First point."]._p.pPr.numPr)
        first_reference = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip().startswith("1.\tFirst paper.")
        )
        self.assertEqual(first_reference.style.name, "Review Reference")
        self.assertIsNone(first_reference._p.pPr.numPr)
        self.assertAlmostEqual(first_reference.style.paragraph_format.left_indent.pt, 36.0)
        self.assertAlmostEqual(first_reference.style.paragraph_format.first_line_indent.pt, -36.0)
        self.assertIsNone(by_text["Second paragraph."].runs[0].font.name)
        self.assertIsNone(by_text["Second paragraph."].runs[0].font.size)

        report = audit_module.audit(output, markdown)
        self.assertEqual(report["layout_profile"], "professional_single")
        self.assertEqual(report["blocking_issues"], [])

    def test_docx_chemvellum_journal_uses_branded_two_column_layout(self) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        module = load_module("md2docx_chemvellum_journal", MD2DOCX)
        audit_module = load_module("audit_docx_chemvellum_journal", AUDIT_DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        markdown = self.root / "chemvellum.md"
        markdown.write_text(
            "# ChemVellum Review\n\n"
            "## Abstract\n\n"
            "A compact evidence-led abstract.\n\n"
            "## 1. Introduction\n\n"
            "Lead paragraph.\n\nSecond paragraph.\n\n"
            "Table 1. Representative catalyst performance.\n\n"
            "| Catalyst | Electrolyte | Operating point | NH_3_ FE (%) | Yield rate | Stability | Ref. |\n"
            "|---|---|---|---|---|---|---|\n"
            "| Cu | 0.1 M KOH | -0.4 V vs RHE | 90 | 1 mmol g^-1^ h^-1^ | 10 h | [1] |\n"
            "| Fe | 0.1 M NaOH | -0.5 V vs RHE | 80 | n.r. | n.r. | [2] |\n\n"
            "Comparison paragraph.\n\n"
            "## References\n\n1. First paper.\n",
            encoding="utf-8",
        )
        output = self.root / "chemvellum.docx"
        module.convert(
            markdown,
            output,
            template,
            layout_profile="chemvellum_journal",
        )
        document = Document(str(output))

        self.assertEqual(document.styles["Review Title"].font.name, "Georgia")
        self.assertEqual(document.styles["Review Body"].font.name, "Cambria")
        self.assertEqual(document.styles["Review Heading 1"].font.name, "Arial")
        body = document.styles["Review Body"].paragraph_format
        self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertAlmostEqual(float(body.line_spacing), 1.04, places=2)
        self.assertAlmostEqual(body.first_line_indent.pt, 9.0, places=2)
        self.assertAlmostEqual(
            document.styles["Review Body Lead"].paragraph_format.first_line_indent.pt,
            9.0,
            places=2,
        )
        self.assertEqual(document.styles["Review Figure Caption"].font.name, "Cambria")
        self.assertEqual(document.styles["Review Figure Caption"].font.size.pt, 8.0)
        self.assertEqual(document.styles["Review Table Header"].font.name, "Arial")
        self.assertEqual(document.styles["Review Table Body"].font.name, "Cambria")

        self.assertGreaterEqual(len(document.sections), 2)
        column_counts = []
        for section in document.sections:
            cols = section._sectPr.find(qn("w:cols"))
            column_counts.append(int(cols.get(qn("w:num")) or 1))
        self.assertEqual(column_counts[0], 1)
        self.assertIn(2, column_counts[1:])
        header_text = "".join(
            paragraph.text
            for section in document.sections
            for paragraph in section.header.paragraphs
        )
        self.assertIn("CHEMVELLUM", header_text)
        abstract = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "Review Abstract"
        )
        self.assertIsNotNone(abstract._p.pPr.find(qn("w:shd")))
        borders = abstract._p.pPr.find(qn("w:pBdr"))
        self.assertIsNotNone(borders)
        self.assertTrue(
            all(borders.find(qn(f"w:{edge}")) is not None for edge in ("top", "left", "right", "bottom"))
        )
        self.assertAlmostEqual(abstract.paragraph_format.right_indent.pt, 13.0, places=2)

        lead = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Lead paragraph.")
        self.assertEqual(lead.style.name, "Review Body Lead")
        table = document.tables[0]
        grid = [
            int(column.get(qn("w:w")) or 0)
            for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        self.assertEqual(grid, [1872, 2664, 1512, 1008, 1872, 893, 547])
        self.assertEqual(
            table.cell(0, 2).paragraphs[0].alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        self.assertIsNotNone(table.cell(2, 0)._tc.get_or_add_tcPr().find(qn("w:shd")))
        comparison = next(
            paragraph for paragraph in document.paragraphs if paragraph.text == "Comparison paragraph."
        )
        self.assertTrue(comparison.paragraph_format.keep_together)

        report = audit_module.audit(output, markdown)
        self.assertEqual(report["layout_profile"], "chemvellum_journal")
        self.assertEqual(report["blocking_issues"], [])

    def test_docx_tables_protect_leading_and_trailing_body_rows(self) -> None:
        from docx import Document

        module = load_module("md2docx_table_orphans", MD2DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        document = Document(str(template))
        module._clear_body(document)
        module._configure_academic_document(document)
        module._add_table_single(
            document,
            ["Method", "Yield"],
            [["A", "80%"], ["B", "75%"], ["C", "70%"], ["D", "65%"]],
        )
        table = document.tables[0]
        for row_index in (0, 1, len(table.rows) - 2):
            self.assertTrue(
                all(
                    paragraph.paragraph_format.keep_with_next
                    for cell in table.rows[row_index].cells
                    for paragraph in cell.paragraphs
                )
            )

    def test_docx_ordered_lists_have_explicit_restart(self) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        module = load_module("md2docx_numbering", MD2DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        document = Document(str(template))
        num_id = module._create_numbering_definition(
            document,
            ordered=True,
        )
        numbering = document.part.numbering_part.element
        target = next(
            item
            for item in numbering.findall(qn("w:num"))
            if item.get(qn("w:numId")) == str(num_id)
        )
        overrides = target.findall(qn("w:lvlOverride"))
        self.assertEqual(len(overrides), 3)
        self.assertTrue(
            all(
                item.find(qn("w:startOverride")).get(qn("w:val")) == "1"
                for item in overrides
            )
        )

    def test_docx_numbering_definitions_have_distinct_word_identities(self) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        module = load_module("md2docx_unique_numbering", MD2DOCX)
        template = REPO / "skills" / "review-export-docx" / "review_template.docx"
        document = Document(str(template))
        module._create_numbering_definition(document, ordered=True)
        module._create_numbering_definition(document, ordered=True, reference=True)
        numbering = document.part.numbering_part.element
        custom_abstracts = numbering.findall(qn("w:abstractNum"))[-2:]
        nsids = [
            item.find(qn("w:nsid")).get(qn("w:val"))
            for item in custom_abstracts
        ]
        templates = [
            item.find(qn("w:tmpl")).get(qn("w:val"))
            for item in custom_abstracts
        ]
        self.assertEqual(len(set(nsids)), 2)
        self.assertEqual(len(set(templates)), 2)
        self.assertEqual(
            custom_abstracts[-1].find(qn("w:multiLevelType")).get(qn("w:val")),
            "singleLevel",
        )
        self.assertEqual(len(custom_abstracts[-1].findall(qn("w:lvl"))), 1)

    def test_docx_list_continuations_stay_inside_their_list_item(self) -> None:
        module = load_module("md2docx_list_continuations", MD2DOCX)
        blocks = module.tokenize(
            "1. First finding begins here\n"
            "   and continues on the next source line.\n"
            "2. Second finding begins here\n"
            "   and also continues.\n"
        )
        list_items = [block for block in blocks if block.kind == "list_item"]
        self.assertEqual(
            [(block.list_number, block.text) for block in list_items],
            [
                (1, "First finding begins here and continues on the next source line."),
                (2, "Second finding begins here and also continues."),
            ],
        )

    def test_docx_plus_prefixed_condition_is_paragraph_continuation(self) -> None:
        module = load_module("md2docx_plus_condition", MD2DOCX)
        blocks = module.tokenize(
            "The electrolyte contained\n+ 1 M KOH [20].\n"
        )
        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0].kind, "paragraph")
        self.assertEqual(blocks[0].text, "The electrolyte contained + 1 M KOH [20].")

    def test_docx_references_keep_their_explicit_numbers(self) -> None:
        module = load_module("md2docx_references", MD2DOCX)
        blocks = module.tokenize(
            "## Conclusions\n\n1. First point.\n2. Second point.\n\n"
            "## References\n\n1. First paper.\n2. Second paper.\n"
        )
        reference_items = [
            block
            for block in blocks
            if block.kind == "list_item" and block.text.endswith("paper.")
        ]
        self.assertEqual(
            [(block.list_number, block.text) for block in reference_items],
            [(1, "First paper."), (2, "Second paper.")],
        )

    def test_render_report_records_artifacts_without_visual_verdict(self) -> None:
        module = load_module("render_docx_lightweight", RENDER)
        docx = self.root / "review.docx"
        pdf = self.root / "review.pdf"
        pages = self.root / "pages"
        report_path = self.root / "render.json"
        docx.write_bytes(b"docx fixture")

        def fake_render(_docx: Path, output_pdf: Path):
            output_pdf.write_bytes(b"%PDF-1.4 fixture")
            return True, {"renderer": "fixture", "available": True, "exit_code": 0}

        def fake_rasterize(_pdf: Path, pages_dir: Path):
            pages_dir.mkdir(parents=True, exist_ok=True)
            page = pages_dir / "page-1.png"
            page.write_bytes(b"png fixture")
            return [str(page)], {"available": True, "exit_code": 0}

        module.render_with_libreoffice = fake_render
        module.render_with_word = fake_render
        module.rasterize = fake_rasterize
        module.pdf_page_count = lambda _pdf: 1
        module.layout_warnings = lambda _pdf, _pages: []
        code = module.run(
            SimpleNamespace(
                input=docx,
                output_pdf=pdf,
                pages_dir=pages,
                report=report_path,
            )
        )
        report = json.loads(report_path.read_text(encoding="utf-8"))
        self.assertEqual(code, 0)
        self.assertTrue(report["pdf_created"])
        self.assertTrue(report["page_images_created"])
        rendered = json.dumps(report).lower()
        self.assertNotIn('"render_status"', rendered)
        self.assertNotIn('"inspection_status"', rendered)
        self.assertNotIn('"verdict"', rendered)

    def test_render_pdf_defaults_to_docx_title_stem(self) -> None:
        module = load_module("render_docx_default_title", RENDER)
        docx = self.root / "Paper Title.docx"
        pages = self.root / "pages"
        report_path = self.root / "render.json"
        docx.write_bytes(b"docx fixture")

        def fake_render(_docx: Path, output_pdf: Path):
            output_pdf.write_bytes(b"%PDF-1.4 fixture")
            return True, {"renderer": "fixture", "available": True, "exit_code": 0}

        module.render_with_libreoffice = fake_render
        module.render_with_word = fake_render
        module.rasterize = lambda _pdf, _pages: ([], {"available": True, "exit_code": 0})
        module.pdf_page_count = lambda _pdf: 1
        code = module.run(
            SimpleNamespace(
                input=docx,
                output_pdf=None,
                pages_dir=pages,
                report=report_path,
            )
        )
        report = json.loads(report_path.read_text(encoding="utf-8"))
        self.assertEqual(code, 1)
        self.assertEqual(Path(report["output_pdf"]).name, "Paper Title.pdf")

    def test_render_uses_a_short_workspace_for_long_windows_paths(self) -> None:
        module = load_module("render_docx_short_workspace", RENDER)
        docx = self.root / "review.docx"
        pdf = self.root / "review.pdf"
        pages = self.root / "pages"
        report_path = self.root / "render.json"
        docx.write_bytes(b"docx fixture")
        seen_paths: list[tuple[Path, Path]] = []

        def fake_render(input_docx: Path, output_pdf: Path):
            seen_paths.append((input_docx, output_pdf))
            output_pdf.write_bytes(b"%PDF-1.4 fixture")
            return True, {"renderer": "fixture", "available": True, "exit_code": 0}

        module.needs_short_working_path = lambda _path: True
        module.render_with_libreoffice = fake_render
        module.render_with_word = fake_render
        module.rasterize = lambda _pdf, _pages: ([], {"available": True, "exit_code": 0})
        module.pdf_page_count = lambda _pdf: 1
        module.layout_warnings = lambda _pdf, _pages: []
        code = module.run(
            SimpleNamespace(
                input=docx,
                output_pdf=pdf,
                pages_dir=pages,
                report=report_path,
            )
        )
        report = json.loads(report_path.read_text(encoding="utf-8"))
        self.assertEqual(code, 1)
        self.assertTrue(pdf.is_file())
        self.assertNotEqual(seen_paths[0][0], docx)
        self.assertNotEqual(seen_paths[0][1], pdf)
        self.assertTrue(report["renderer_attempts"][0]["used_short_working_path"])


if __name__ == "__main__":
    unittest.main()
